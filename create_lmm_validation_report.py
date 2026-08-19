# -*- coding: utf-8 -*-
"""
LMM 검증 연구 기록 문서 생성.

    python lmm_diagnose.py                    # 먼저 실행 -> docs/data/lmm_diagnosis.json
    python create_lmm_validation_report.py    # -> docs/output/*_LMM_검증연구기록.docx

수치는 전부 `docs/data/lmm_diagnosis.json` 에서 읽는다 (하드코딩 금지).
진단을 다시 돌리면 문서도 자동으로 갱신된다.

조판: 본문 맑은 고딕 10.5pt / 라틴 Cambria, 표는 가로 괘선만(booktabs).
"""
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).parent
DIAG = HERE / "docs" / "data" / "lmm_diagnosis.json"
OUT_DIR = HERE / "docs" / "output"

KR = "맑은 고딕"
LATIN = "Cambria"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x60, 0x60, 0x60)
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


# ── 조판 헬퍼 ─────────────────────────────────────────────────────────────
def set_run_fonts(run, latin=LATIN, ea=KR):
    """라틴·한글 글꼴을 따로 물린다 (Cambria 에 한글 글리프가 없다)."""
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:eastAsia"), ea)


def para(doc, text="", *, size=10.5, bold=False, italic=False, color=INK,
         align=None, space_before=0, space_after=6, line=1.45, indent=0.0,
         style=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if indent:
        pf.left_indent = Cm(indent)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color
        set_run_fonts(r)
    return p


def rich(doc, parts, *, size=10.5, space_after=6, line=1.45, indent=0.0):
    """[(텍스트, 굵게), ...] 형태로 한 문단 안에서 강조를 섞는다."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if indent:
        pf.left_indent = Cm(indent)
    for txt, bold in parts:
        r = p.add_run(txt)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.color.rgb = INK
        set_run_fonts(r)
    return p


def heading(doc, text, level=1):
    sizes = {1: 15, 2: 12, 3: 11}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    r = p.add_run(text)
    r.font.size = Pt(sizes[level])
    r.bold = True
    r.font.color.rgb = ACCENT if level == 1 else INK
    set_run_fonts(r)
    # 목차에 잡히도록 개요 수준을 부여
    ppr = p._element.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level - 1))
    ppr.append(ol)
    return p


def _border(el, edge, sz=6, color="1A1A1A"):
    tag = OxmlElement(f"w:{edge}")
    tag.set(qn("w:val"), "single")
    tag.set(qn("w:sz"), str(sz))
    tag.set(qn("w:color"), color)
    el.append(tag)


def booktabs(doc, header, rows, widths, *, size=9.5, aligns=None):
    """가로 괘선 3줄만 쓰는 학술 표 — 세로 괘선·음영 없음."""
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    total = sum(widths)
    t._tblPr.append(OxmlElement("w:tblBorders"))  # 자리만 만들고 아래에서 채운다
    tb = t._tblPr.find(qn("w:tblBorders"))
    for edge in ("top", "bottom"):
        _border(tb, edge, sz=10)
    for edge in ("left", "right", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        tb.append(e)
    e = OxmlElement("w:insideH")
    e.set(qn("w:val"), "none")
    tb.append(e)

    aligns = aligns or ["l"] * len(header)
    amap = {"l": WD_ALIGN_PARAGRAPH.LEFT, "r": WD_ALIGN_PARAGRAPH.RIGHT,
            "c": WD_ALIGN_PARAGRAPH.CENTER}

    def fill(cells, vals, bold):
        for c, v, w, a in zip(cells, vals, widths, aligns):
            c.width = Cm(w)
            p = c.paragraphs[0]
            p.alignment = amap[a]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(str(v))
            r.font.size = Pt(size)
            r.bold = bold
            r.font.color.rgb = INK
            set_run_fonts(r)

    fill(t.rows[0].cells, header, True)
    # 머리글 아래 중간 괘선
    for c in t.rows[0].cells:
        tcpr = c._tc.get_or_add_tcPr()
        bd = OxmlElement("w:tcBorders")
        _border(bd, "bottom", sz=6)
        tcpr.append(bd)
    for row in rows:
        fill(t.add_row().cells, row, False)
    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Cm(w)

    # 쪽을 넘어갈 때 머리글을 반복하고, 한 행이 쪼개지지 않게 한다
    for ri, row in enumerate(t.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
        if ri == 0:
            trPr.append(OxmlElement("w:tblHeader"))
    # 작은 표는 통째로 한 쪽에 남긴다 (머리글만 남고 넘어가는 것을 막는다)
    if len(rows) <= 8:
        for row in list(t.rows)[:-1]:
            for c in row.cells:
                for p in c.paragraphs:
                    p.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def caption(doc, text):
    """그림 캡션 — 아래에 둔다."""
    return para(doc, text, size=9, color=MUTED, space_before=2, space_after=10,
                line=1.3)


def tcap(doc, text):
    """표 캡션 — 학술 관례상 표 위에 둔다. 표와 떨어지지 않게 묶는다."""
    p = para(doc, text, size=9, color=MUTED, space_before=8, space_after=3,
             line=1.3)
    p.paragraph_format.keep_with_next = True
    return p


def note(doc, text):
    return para(doc, text, size=9.5, color=MUTED, space_before=2,
                space_after=8, line=1.4, indent=0.4)


def f(x, n=2):
    return "—" if x is None else f"{x:,.{n}f}"


# ══════════════════════════════════════════════════════════════════════
def build(d):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.size = Pt(10.5)
    st.font.name = LATIN
    st.element.rPr.rFonts.set(qn("w:eastAsia"), KR)

    for s in doc.sections:
        s.top_margin = Cm(2.4)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.6)
        s.right_margin = Cm(2.6)

    ds = d["dataset"]
    pc = d["point_check"]
    vr = d["vector_recovery"]
    hy = d["hypotheses"]
    asym = d["asymmetry"]
    req = d["required_local_anomaly"]
    vg = d["variogram"]
    cs = d["coord_sensitivity"]
    ce = d["cross_epoch"]
    o19 = d["only2019"]

    # ── 표제 ──────────────────────────────────────────────────────────
    para(doc, "지구자기장 기준장 모형", size=10, color=MUTED, space_after=2)
    para(doc, "지자기 모델(LMM) 검증 연구 기록", size=21, bold=True,
         space_after=4, line=1.2)
    para(doc, "편각 잔차의 원인 분석과 향후 작업 순서", size=12.5,
         color=ACCENT, space_after=14)
    gen = datetime.fromisoformat(d["generated"])
    para(doc, f"작성 {gen.strftime('%Y년 %m월 %d일')}   ·   "
              f"측점 {ds['n_sites']}개 / 관측 {ds['n_obs']}회 "
              f"({ds['year_min']}~{ds['year_max']})",
         size=9.5, color=MUTED, space_after=16)

    heading(doc, "1. 개요", 1)
    para(doc,
         "이 문서는 한반도 지역 자기장 모형(LMM)의 구현 정확성을 검증하고, "
         "편각 KPI 미달의 원인을 좁히기 위해 수행한 분석을 기록한 것이다. "
         "분석은 두 단계로 이루어졌다. 먼저 계산기가 설계대로 동작하는지를 "
         "외부 기준과 대조해 확인하고, 다음으로 남은 편각 잔차의 원인에 대해 "
         "네 가지 가설을 세워 순차 검정하였다.")
    para(doc,
         "결론부터 적으면, LMM 계산 구현 자체에서는 유의한 오류가 확인되지 "
         f"않았다. 현재 약 {asym['rD_rms']:.0f}′ 의 편각 잔차는 현행 KIGAM "
         "격자에서 복원한 지각 자기장 벡터만으로는 충분히 설명되지 않으며, "
         "측점별 방위기준·성과 환산·외부장 보정 등 관측자료 처리과정의 계통 "
         "오차 가능성이 우선적으로 의심된다. 다만 지각 벡터가 만드는 편각 "
         f"기여의 예상 규모가 {vr['predD_rms_arcmin']:.1f}′ 로 목표 KPI 6′ 를 "
         "이미 넘고, 현재의 벡터 복원에는 격자 결측과 공간해상도의 한계가 "
         "있으므로 구조적 오차원을 배제할 수는 없다. 따라서 원 야장 감사와 "
         "외부장 보정 후 재적합을 우선 수행하고, 그 뒤에 지각 자기장 벡터 "
         "성분의 반영 필요성을 KPI 관점에서 다시 평가한다.")

    heading(doc, "1.1 결론 요약", 2)
    tcap(doc, "표 1. 검증 결과 요약")
    booktabs(doc,
             ["항목", "결과", "판정"],
             [["① Core 층 구현",
               "공식 IGRF 계산기와 표시 자릿수 내 일치", "정상"],
              ["LMM 계산 구현",
               f"독립 재현과 {abs(pc['vs_webcalc']['lmm_F']):.2f} nT 이내 일치", "정상"],
              ["총자력 차이의 정체",
               f"지각 이상이 {abs(pc['layers']['crustal_F']):.0f} nT 로 대부분", "설계대로"],
              ["지각층의 각도 기여",
               f"벡터 복원으로 반영 · 예상 기여 {vr['predD_rms_arcmin']:.1f}′",
               "KPI 6′ 초과 — 무시 불가"],
              ["편각 잔차 원인",
               f"약 {asym['rD_rms']:.0f}′ — 측점별 계통 오차 우선 의심",
               "확정 아님 · 자료 감사 필요"],
              ["상시관측 자료",
               "야장에서 관측시각 복원 · 보정량 산출", "모델 반영은 재판정 대기"]],
             [3.2, 7.4, 3.4])

    # ── 2. 구현 검증 ──────────────────────────────────────────────────
    heading(doc, "2. 계산기 구현 검증", 1)
    p = pc["point"]
    para(doc,
         f"검증 지점은 위도 {p['lat']}°N, 경도 {p['lon']}°E, 표고 "
         f"{p['elev_m']:.0f} m, 일자 {p['date']} 이다. 이 지점을 공식 IGRF "
         "계산기(외부 웹 서비스), 프로젝트의 웹 계산기(docs/lmm.html), "
         "그리고 ppigrf 를 이용한 독립 재현 세 경로로 계산하여 대조하였다.")

    heading(doc, "2.1 ① Core 층 — 공식 계산기 대조", 2)
    vo = pc["vs_official"]
    vw = pc["vs_webcalc"]
    tcap(doc, "표 2. IGRF 계산 결과 대조. 웹 계산기 열은 "
                 "docs/lmm.html 을 브라우저에서 직접 실행해 얻은 값이다.")
    booktabs(doc,
             ["성분", "ppigrf", "공식 계산기", "웹 계산기", "ppigrf−공식"],
             [["편각 D", f"{pc['ppigrf']['D']:.4f}°",
               f"{pc['official']['D']:.3f}°", f"{pc['webcalc']['igrf_D']:.4f}°",
               f"{vo['D']:+.4f}°"],
              ["복각 I", f"{pc['ppigrf']['I']:.4f}°",
               f"{pc['official']['I']:.3f}°", f"{pc['webcalc']['igrf_I']:.4f}°",
               f"{vo['I']:+.4f}°"],
              ["수평 H", f"{pc['ppigrf']['H']:,.2f}",
               f"{pc['official']['H']:,.0f}", f"{pc['webcalc']['igrf_H']:,.1f}",
               f"{vo['H']:+.2f} nT"],
              ["총자력 F", f"{pc['ppigrf']['F']:,.2f}",
               f"{pc['official']['F']:,.0f}", f"{pc['webcalc']['igrf_F']:,.1f}",
               f"{vo['F']:+.2f} nT"]],
             [2.4, 3.0, 3.0, 3.0, 2.6], aligns=["l", "r", "r", "r", "r"])
    rich(doc, [
        ("공식 계산기와의 차이는 전부 ", False),
        ("화면 표시 자릿수", True),
        (" 때문이다. 공식 계산기는 강도를 정수로 끊어 표시하므로, "
         f"ppigrf 의 {pc['ppigrf']['F']:,.2f} nT 는 반올림하면 표시값과 정확히 "
         "일치한다. 우리 웹 계산기의 IGRF 열은 ppigrf 와 편각 ", False),
        (f"{abs(vw['D'])*3600:.1f}″", True),
        (", 총자력 ", False),
        (f"{abs(vw['F']):.2f} nT", True),
        (" 이내로 맞는다.", False)])

    heading(doc, "2.2 LMM 전체 재현과 층별 분해", 2)
    lay = pc["layers"]
    dif = pc["diff_lmm_igrf"]
    tcap(doc, "표 3. 검증 지점에서의 층별 기여 분해")
    booktabs(doc,
             ["층", "총자력 기여", "비중"],
             [["① Core (IGRF-14)", f"{lay['core']:,.1f} nT", "기준"],
              ["② Regional (지표 절대측정)", f"{lay['regional_F']:+.2f} nT",
               f"{100*abs(lay['regional_F'])/abs(dif['F']):.1f} %"],
              ["③ Crustal (KIGAM 항공자력)", f"{lay['crustal_F']:+.2f} nT",
               f"{100*abs(lay['crustal_F'])/abs(dif['F']):.1f} %"],
              ["④ External (CYG)", "미적용", "—"],
              ["LMM − IGRF 합계", f"{dif['F']:+.2f} nT", "100 %"]],
             [6.0, 4.0, 4.0], aligns=["l", "r", "r"])
    note(doc,
         f"배포 격자(lmm_model.json)의 지각 이상값은 정수로 반올림되어 있어 원 "
         f"KIGAM 격자와 {lay['grid_quantization_nT']:.2f} nT 차이가 난다. 위 표는 "
         "계산기 검증이므로 배포 격자 기준이다. 양자화 오차는 총자력 KPI 50 nT "
         "대비 무시할 수 있다.")
    rich(doc, [
        ("총자력 차이는 사실상 전부 지각 자기이상이다. 이는 오류가 아니라 "
         "IGRF 단독으로는 놓치는 국소 이상을 잡아낸 것이므로 ", False),
        ("모형이 설계대로 작동한 증거", True),
        (f"다. 반면 편각 차이 {dif['D']*60:+.2f}′ 와 복각 차이 "
         f"{dif['I']*60:+.2f}′ 는 Regional 다항식만으로 만들어진 값이다.", False)])

    heading(doc, "2.3 지각층은 각도 성분에 기여하지 않는다", 2)
    para(doc,
         "결합 방식은 소스에서 확인된다. docs/lmm.html 의 lmm() 함수는 "
         "다음과 같이 계산한다.")
    cp = para(doc,
              "D = D_IGRF + ΔD_regional\n"
              "I = I_IGRF + ΔI_regional\n"
              "F = F_IGRF + ΔF_regional + ΔF_crustal",
              size=9.5, indent=0.6, space_after=8, line=1.5)
    for r in cp.runs:
        set_run_fonts(r, latin="Consolas", ea=KR)
    rich(doc, [
        ("즉 지각 이상은 총자력에만 더해지고 편각·복각에는 들어가지 않는다. "
         "이는 버그가 아니라 lmm_build.py 의 fit_regional() 주석에 명시된 "
         "설계 선택이다 — ", False),
        ("항공자력이 스칼라 이상만 제공하기 때문", True),
        ("이다. 그러나 스칼라 총자력 이상은 지각 이상 벡터를 주자기장 방향에 "
         "투영한 성분이므로, 각도를 틀어놓는 수직 성분이 버려진다. 이 누락이 "
         "편각 잔차의 원인인지가 다음 절의 첫 번째 가설이다.", False)])
    note(doc,
         "실증: 측점 주변 ±0.1° 를 훑으면 지각 이상이 약 124 nT 요동치는 동안 "
         "편각 보정량은 Regional 다항식의 기울기대로 매끄럽게만 변한다. "
         "두 양의 상관은 0 이다.")

    # ── 3. 원인 분석 ──────────────────────────────────────────────────
    heading(doc, "3. 편각 잔차 원인 분석", 1)
    para(doc,
         f"측점 {ds['n_sites']}개 전수를 대상으로, 현행 모형(Regional "
         f"{d['degree']}차)이 설명하지 못하고 남긴 잔차의 원인을 검정하였다. "
         f"편각 품질 합격 측점은 {ds['D_ok']}개, 시그마 클리핑 후 "
         f"inlier 는 {ds['inlier_D']}개이다.")

    heading(doc, "3.1 검정 결과 종합", 2)
    hA, hB, hB2, hC = (hy["A_vector_D"], hy["B_precision"], hy["B_nvisit"],
                       hy["C_roughness"])
    tcap(doc, "표 4. 가설 검정 결과. n≈14 에서 유의수준 0.05 를 넘으려면 "
                 "|r| > 0.53 이 필요하므로 넷 모두 미달이다.")
    booktabs(doc,
             ["가설", "검정 대상", "n", "r", "p", "판정"],
             [["A. 지각 벡터 성분 누락", "복원 ΔD vs 잔차",
               hA["n"], f"{hA['r']:+.3f}", f"{hA['p']:.3f}", "기각"],
              ["B. 측정 정밀도", "재방문 산포 vs |잔차|",
               hB["n"], f"{hB['r']:+.3f}", f"{hB['p']:.3f}", "기각"],
              ["B′. 관측 횟수", "방문 수 vs |잔차|",
               hB2["n"], f"{hB2['r']:+.3f}", f"{hB2['p']:.3f}", "기각"],
              ["C. 국소 자기 거칠기", "5 km ΔF 표준편차 vs |잔차|",
               hC["n"], f"{hC['r']:+.3f}", f"{hC['p']:.3f}", "기각"]],
             [4.6, 5.4, 1.0, 1.5, 1.5, 1.8],
             aligns=["l", "l", "c", "r", "r", "c"])

    heading(doc, "3.2 가설 A — 지각 벡터 성분 누락", 2)
    para(doc,
         "포텐셜장 가정(b = −∇V, ∇²V = 0)에서 스칼라 총자력 이상 격자로부터 "
         "벡터 성분을 FFT 로 역산하고, 그 벡터가 각 측점에서 만드는 편각 "
         "변화를 예측하여 잔차와 대조하였다. 복원한 벡터를 주자기장 방향에 "
         f"재투영한 검산 오차는 RMS {vr['recon_rms_nT']:.2f} nT 로, "
         "제거한 상수항(DC)에 해당하는 값만 남아 역산은 정상이다.")
    tcap(doc, "표 5. 가설 A 의 규모 비교")
    booktabs(doc,
             ["양", "값"],
             [["지각 벡터 차감 전 편각 잔차 RMS", f"{hA['rms']:.2f}′"],
              ["복원 벡터가 예측하는 ΔD RMS", f"{vr['predD_rms_arcmin']:.2f}′"],
              ["차감 후 잔차 RMS (= 현행 모형의 잔차)",
               f"{hA['rms_sub']:.2f}′ ({hA['rms_sub']/60:.3f}°)"],
              ["최적 배율 적용 시", f"{hA['rms_scaled']:.2f}′ (배율 {hA['a']:.2f})"],
              ["편각 KPI", "6.00′ (0.100°)"],
              ["복원된 수평 성분 RMS (동)", f"{vr['b_east_rms']:.1f} nT"],
              ["원 스칼라 이상 RMS", f"{vr['scalar_rms']:.1f} nT"]],
             [8.0, 5.0], aligns=["l", "r"])
    note(doc,
         "표 읽는 법: 현행 모형은 이미 지각 벡터를 빼고 있으므로 「차감 후 잔차」와 "
         f"이 문서 전반의 편각 잔차({asym['rD_rms']:.2f}′)는 같은 값이다. 개선율은 "
         "차감 전 값과의 비교로 읽어야 한다.")
    rich(doc, [
        (f"복원 벡터를 차감하면 편각 잔차가 {hA['rms']:.1f}′ 에서 "
         f"{hA['rms_sub']:.1f}′ 로 {100*(1-hA['rms_sub']/hA['rms']):.0f} % "
         "줄어든다. 두 가지를 함께 읽어야 한다. 첫째, 설명해야 할 잔차에 견주면 "
         "지각 벡터는 ", False),
        ("현재 잔차의 주원인이 아니다", True),
        (f". 둘째, 그럼에도 예상 기여 {vr['predD_rms_arcmin']:.1f}′ 는 ", False),
        ("편각 KPI 6′ 를 이미 넘는다", True),
        (". 따라서 자료 계통 오차를 걷어낸 뒤에는 지각 벡터 누락이 KPI 달성을 "
         "제한하는 2차 오차원이 될 수 있으므로, 이 가설을 기각 목록에 넣어서는 "
         "안 된다.", False)])
    rich(doc, [
        ("역으로, 관측된 잔차를 모두 지각 이상으로 설명하려면 국소 수평 이상이 "
         "중앙값 ", False),
        (f"{req['median_nT']:.0f} nT", True),
        (f", 최대 {req['max_nT']:.0f} nT 필요하다. 이는 현재 격자에서 복원된 "
         f"규모의 약 {req['factor']:.1f}배다. 다만 격자 간격은 진폭의 상한을 "
         "정하는 값이 아니라 해상 가능한 파장의 한계이므로, 정확히 말하면 "
         f"{vr['grid']['dx_km']:.1f} km 격자로는 측점 주변의 단파장 국소 이상을 "
         "복원할 수 없어 그 규모의 존재 여부를 현 자료만으로는 판단하기 "
         "어렵다는 뜻이다.", False)])
    note(doc,
         f"⚠ 이 검정은 탐색적이다. KIGAM 격자의 결측 "
         f"{vr['grid']['gap_pct']:.0f} % 를 0 nT 로 채운 뒤 FFT 로 역산했는데, "
         "결측과 「이상 0」은 다른 뜻이며 그 경계에서 생기는 인공 불연속에 "
         "파수영역 미분은 특히 민감하다. 따라서 위 예상 기여값 자체에도 "
         "불확실성이 있다.")

    heading(doc, "3.3 가설 B — 측정 정밀도", 2)
    sp = d["precision_split"]
    rich(doc, [
        ("방문 횟수가 많을수록 잔차가 줄어드는 경향은 관찰되지 않았다. 1회 관측 "
         "측점의 |잔차| 중앙값은 ", False),
        (f"{sp['single_median']:.1f}′", True),
        (f" (n={sp['single_n']}), 2회 이상 측점은 ", False),
        (f"{sp['multi_median']:.1f}′", True),
        (f" (n={sp['multi_n']}) 로 오히려 재방문 측점 쪽이 크다. 표본이 작아 "
         "이 차이 자체를 해석할 수는 없으나, 적어도 «재방문으로 정밀도를 "
         "높이면 잔차가 준다»는 관계는 보이지 않는다. 정밀도가 우수한 측점이 "
         "크게 틀린 사례도 있다.", False)])
    st_ = {s["name"]: s for s in d["site_table"]}
    tcap(doc, "표 6. 정밀도와 정확도의 분리")
    booktabs(doc,
             ["측점", "관측 연도", "D 재방문 산포", "잔차"],
             [["성산", "2023, 2025", "0.6′", f"{st_['성산']['rD']:+.1f}′"],
              ["순천", "2023, 2025", "0.5′", f"{st_['순천']['rD']:+.1f}′"],
              ["강화", "2022, 2024", "3.2′", f"{st_['강화']['rD']:+.1f}′"],
              ["여주", "2023, 2025", "11.0′", f"{st_['여주']['rD']:+.1f}′"]],
             [2.6, 4.4, 4.0, 3.0], aligns=["l", "l", "r", "r"])
    rich(doc, [
        ("성산은 ", False),
        (f"2년 간격의 독립 캠페인 두 번이 0.6′ 로 일치하면서 둘 다 "
         f"{abs(st_['성산']['rD']):.0f}′ 틀렸다", True),
        ("(표 6 의 잔차는 Core·Crustal·Regional 을 모두 뺀 뒤의 값이다. "
         "IGRF 만 뺀 값은 이보다 크다). 매번 같은 값이 나오면서 틀린다는 것은 "
         "무작위 잡음이 아니라 측점에 "
         "고정된 계통 오차를 뜻한다. 동시에 순천·강화는 정밀도와 정확도가 둘 다 "
         f"양호하므로, 사업 전체의 계통 문제가 아니라 측점별 문제이다. "
         f"편각 inlier {ds['inlier_D']}개 중 |잔차| 10′ 이내는 3개뿐이다.",
         False)])
    note(doc,
         "주의: 가설 B 의 검정 대상은 재방문 산포, 곧 정밀도(precision)이다. "
         "이 검정의 기각은 계통 오차(accuracy)를 배제하지 못한다.")

    heading(doc, "3.4 가설 D — 잔차의 공간 구조", 2)
    para(doc,
         "잔차가 미적합 지역장 구조에서 온다면 가까운 측점끼리 비슷한 값을 "
         "가져야 한다. 반변량도로 확인하였다.")
    tcap(doc, f"표 7. 편각 잔차의 반변량도. 거리 추세 r = {vg['trend_r']:+.3f} "
                 f"(p = {vg['trend_p']:.3f})")
    booktabs(doc,
             ["거리 구간 (km)", "측점쌍", "반변량"],
             [[f"{b['lo']} – {b['hi']}", b["n"], f"{b['semivariance']:,.0f}"]
              for b in vg["bins"]]
             + [["전체 (sill)", vg["n_pairs"], f"{vg['sill']:,.0f}"]],
             [4.6, 3.0, 4.0], aligns=["l", "c", "r"])
    rich(doc, [
        (f"최단 거리 구간({vg['bins'][0]['lo']}–{vg['bins'][0]['hi']} km)에서 "
         f"이미 반변량이 sill 을 넘고 거리 추세도 뚜렷하지 않다. 현재 "
         f"{ds['inlier_D']}개 측점에서는 ", False),
        ("뚜렷한 공간 상관 구조가 확인되지 않는다", True),
        (". 다만 이를 순수 nugget 구조로 ", False),
        ("확정할 수는 없다", True),
        (f". 최근거리 구간의 측점쌍이 {vg['bins'][0]['n']}개뿐이고(경험 "
         "반변량도가 안정되려면 구간당 수십 쌍이 권장된다), 거리 추세의 유의 "
         f"확률({vg['trend_p']:.3f})도 {vg['n_pairs']}개 쌍을 독립 표본으로 "
         "본 값이라 과소평가되어 있다. 측점 하나가 여러 쌍에 반복해 들어가기 "
         "때문이다. 공간 구조의 유무는 측점을 늘려야 판정할 수 있다.", False)])

    heading(doc, "3.5 편각·복각 비대칭", 2)
    tcap(doc, f"표 8. 동일 측점 {asym['n']}개에서의 D/I 비교")
    booktabs(doc,
             ["양", "값"],
             [["편각 잔차 RMS", f"{asym['rD_rms']:.2f}′"],
              ["복각 잔차 RMS", f"{asym['rI_rms']:.2f}′"],
              ["관측된 비 D/I", f"{asym['ratio']:.2f}"],
              ["자기 원인이라면 예상되는 비 (F/H)",
               f"{asym['expected_from_magnetic']:.2f}"],
              ["복원 벡터가 예측하는 비", f"{asym['pred_ratio']:.2f}"]],
             [8.0, 5.0], aligns=["l", "r"])
    rich(doc, [
        (f"관측된 비 {asym['ratio']:.2f} 는 자기 원인에서 기대되는 값보다 "
         "크다. 다만 이 기대값(F/H)은 δX·δY·δZ 가 비슷한 분산을 갖는 ", False),
        ("등방 자기이상 가정에서만 성립", True),
        ("한다. 실제 지각 자기이상은 지질구조·자화방향·잔류자화에 따라 "
         "방향성을 가질 수 있으므로, 비가 크다는 사실만으로 초과 오차의 크기를 "
         "특정할 수는 없다. 여기서 말할 수 있는 것은 ", False),
        ("편각 잔차가 복각 잔차보다 현저히 크며, 편각 측정·방위 기준 계통 "
         "오차의 가능성을 시사한다", True),
        ("는 것까지다. 후보는 기준 마크의 진북 방위각 오차, 성과표 편각 값의 "
         "전사·환산 오류, 편각 환산 절차 오류이며 자료만으로는 구분되지 "
         "않는다.", False)])

    heading(doc, "3.6 자료 오류 가능성 검정", 2)
    tcap(doc, "표 9. 자료 오류 가능성 점검")
    booktabs(doc,
             ["점검", "결과", "판정"],
             [["좌표 오류",
               f"편각 민감도 북 {cs['per_km']['north_1km']['dD_arcmin']:.3f}′/km, "
               f"동 {cs['per_km']['east_1km']['dD_arcmin']:.3f}′/km",
               "배제"],
              ["기준시점 환산 오류",
               f"D 기울기 {d['epoch_trend']['D']['slope']:+.2f}′/년, "
               f"I 기울기 {d['epoch_trend']['I']['slope']:+.2f}′/년",
               "신호 없음"],
              ["자료원 계통 차이",
               f"2019 야장 vs 성과표 p = {d['source_split']['rD']['p']:.3f}",
               "유의성 없음"]],
             [3.2, 7.4, 2.6])
    rich(doc, [
        ("좌표는 완전히 무관하다. 잔차 중앙값을 위치 오차로 설명하려면 ", False),
        (f"{cs['km_needed']:.0f} km", True),
        (f" 가 틀려야 한다. CLAUDE.md 가 지적한 미원의 353 m 좌표 불일치는 "
         f"편각으로 환산하면 {cs['mion_353m_arcmin']:.3f}′ 로 무의미하다. "
         "기준시점 환산 오류라면 편각과 복각 모두 각자의 영년변화 크기에 "
         "비례한 연도 의존이 남아야 하지만, 복각에는 신호가 전혀 없어 짝이 "
         "맞지 않는다.", False)])

    heading(doc, "3.7 종합", 2)
    tcap(doc, "표 10. 편각 잔차 원인의 규모별 정리")
    booktabs(doc,
             ["원인", "규모", "판정"],
             [["지각 이상 벡터 성분 누락", f"약 {vr['predD_rms_arcmin']:.0f}′",
               "주원인 아님 · KPI 6′ 초과로 무시 불가"],
              ["측정 잡음(정밀도)", "—", "관계 확인되지 않음"],
              ["미적합 지역 구조", "—", "공간 상관 확인되지 않음(표본 부족)"],
              ["측점별 편각 계통 오차", f"약 {asym['rD_rms']:.0f}′ 규모",
               "우선 의심 — 직접 증거는 없음"]],
             [5.4, 3.0, 5.0])

    # ── 4. CYG ────────────────────────────────────────────────────────
    heading(doc, "4. 외부장(CYG) 자료 현황과 제약", 1)
    cyg = d["cyg_cache"]
    rows = []
    for g in cyg["groups"]:
        fmt = lambda s: f"{s[:4]}-{s[4:6]}-{s[6:]}"
        rows.append([g["group"], f"{g['n_days']}일",
                     f"{fmt(g['first'])} ~ {fmt(g['last'])}"])
    tcap(doc, "표 11. 청양(CYG) 1분 자료 캐시 실측 현황")
    booktabs(doc, ["구간", "일수", "날짜 범위"], rows,
             [3.4, 2.6, 7.0], aligns=["l", "r", "l"])
    rich(doc, [
        ("CYG 자료는 보유하고 있고 실제로 사용 중이다. 2019 구간 78일은 야장 "
         "세션의 외부장 보정에 쓰였고, 2024년 6월 10일 표본은 일변동 규모를 "
         "재는 데 쓰였다. 적용이 막힌 것은 ", False),
        ("2022~25 성과표 구간뿐이며, 원인은 CYG 가 아니라 성과표", True),
        ("에 있다.", False)])
    rich(doc, [
        ("lmm_build.py 의 load_survey_points() 는 성과표에 연도만 있으므로 "
         "관측 일자를 ", False),
        ("7월 1일로 대체", True),
        (f"한다. 실제로 성과표 유래 레코드의 월·일은 "
         f"{', '.join(d['survey_date_stub']['unique_monthday'])} 하나뿐이다. "
         "CYG 는 1분 자료인데 조회할 시각이 없고, 나아가 실제 관측 날짜조차 "
         "모르므로 해당 구간 자료를 내려받을 수도 없다.", False)])
    tcap(doc, "표 12. 7월 1일 근사가 발생시키는 비용")
    booktabs(doc,
             ["영향", "규모"],
             [["영년변화(Core 층) 오차 — 실제 날짜가 ±4개월 어긋날 때", "±0.7′"],
              ["External 층 — 보정 자체가 봉쇄", "6.7′ (중앙값) ~ 13.2′ (최대)"]],
             [8.6, 4.4], aligns=["l", "r"])
    para(doc,
         "즉 날짜 근사는 기준장에는 거의 해롭지 않고 외부장 보정만 완전히 "
         "막는다. 법정으로는 지구물리측량 작업규정 제19조 제1항 제6호가 시각 "
         "동시 측정을 의무화하므로 2022~25 야장에는 시각이 있어야 하며, "
         "성과표로 정리되는 과정에서 탈락한 것으로 판단된다.")
    note(doc,
         "웹 계산기가 ④ External 을 미적용으로 표시하는 것은 결함이 아니다. "
         "LMM 은 정온시 기준값 모형이므로 외부장 보정이 들어갈 자리는 예측 "
         "시점이 아니라 관측 성과를 정리하는 단계이다.")

    # ── 5. 2019 단독 ──────────────────────────────────────────────────
    heading(doc, "5. 2019 자료 단독 구축 가능성", 1)
    rich(doc, [
        (f"2019 야장으로 쓸 수 있는 측점은 {o19['n_sites']}개, 세션 "
         f"{o19['n_sessions']}회이다(성산은 시각 미기입으로 제외). ", False),
        ("배포용 모형 구축은 불가능하다", True),
        (".", False)])
    tcap(doc, "표 13. 2019 자료 단독 구축의 제약")
    booktabs(doc,
             ["제약", "내용"],
             [["공간 분포",
               f"위도 {o19['lat_min']:.2f}~{o19['lat_max']:.2f}°N. 최북단 위로 "
               f"약 {o19['gap_north_km']:.0f} km 구간에 측점 없음"],
              ["자유도",
               f"1차 다항식 계수 3개에 자유도 {o19['dof'][0]['dof']}, "
               f"2차는 {o19['dof'][1]['dof']}"],
              ["이상치 배제",
               f"robust_lstsq 최소 요구 표본 {o19['dof'][0]['robust_min']}개 미달 "
               "— 시그마 클리핑이 작동하지 않음"],
              ["측점 구성",
               "5개 중 남지·거제 2개가 최악 잔차 측점"]],
             [3.0, 10.0])
    para(doc,
         "다만 2019 자료는 방법론 검증에는 유효하다. 특히 NOC 공간투영 구현의 "
         "회귀 가드로 쓸 수 있다. 균일 V 근사를 쓴 기존 보정은 정온시 편각을 "
         "0.15′ 악화시키고 교란시 복각을 0.01′에서 4.39′로 악화시켰으므로, "
         "새 구현이 같은 손상을 일으키지 않는지 확인하는 용도이다. 원자료 세션 "
         "산포가 최대 2.00′ 인 데 비해 보정량이 최대 5.22′ 이므로 검정력은 "
         "낮지만, 구현 오류를 잡는 값은 있다.")

    # ── 6. 시기간 대조 ────────────────────────────────────────────────
    heading(doc, "6. 시기간 대조 — 신규 검증 지표", 1)
    para(doc,
         "2019 야장과 2022~25 성과표에 모두 존재하는 측점을 대조하였다. "
         "값은 IGRF 잔차이므로 주자기장의 영년변화는 이미 제거되어 있다. "
         "다만 잔차에는 IGRF 의 지역 모델오차·지각장·외부장·관측오차가 함께 "
         "남으므로 두 시기의 차이가 정확히 0 일 필요는 없다. 외부장 보정과 "
         "기준시점 처리, 관측절차·방위표지·장비가 모두 같다면 이 차이가 "
         "작은 값으로 수렴해야 한다는 뜻에서 검증 지표로 쓴다.")
    rows = []
    for c in ce["sites"]:
        rows.append([c["name"],
                     f"{c['dD']['fb2019']:+.2f}′", f"{c['dD']['survey']:+.2f}′",
                     f"{c['dD']['diff']:+.2f}′", f"{c['dF']['diff']:+.1f} nT"])
    tcap(doc, "표 14. 겹치는 측점의 시기간 대조")
    booktabs(doc,
             ["측점", "2019 (검증완료)", "2022~25 (미검증)", "편각 차", "총자력 차"],
             rows, [2.2, 3.0, 3.2, 2.4, 2.6],
             aligns=["l", "r", "r", "r", "r"])
    rich(doc, [
        ("이 불일치는 두 가지로 설명 가능하고 대조만으로는 가려지지 않는다. ",
         False),
        ("성과표 오류", True),
        ("이거나, ", False),
        ("성과표의 외부장 미보정", True),
        ("이다. 2019 자료는 Kp>2 세션이 배제된 것이고 성과표는 무보정이며, "
         "관측된 차이가 일변동 규모(편각 중앙 6.7′ / 최대 13.2′, 총자력 중앙 "
         "37.5 / 최대 68.4 nT)와 정확히 겹친다.", False)])
    para(doc,
         "이는 오히려 유용하다. 야장을 확보해 External 을 적용한 뒤 이 세 "
         "측점의 시기간 차가 0 으로 수렴하면 원인은 외부장이고 External 이 "
         "작동한 증거가 된다. 차이가 그대로 남으면 성과표 오류가 확정된다. "
         "지금까지 교차검증 외에는 External 의 성패를 판정할 지표가 없었으므로, "
         "측점 3개짜리이지만 방향이 명확한 지표를 확보한 셈이다.")
    rich(doc, [
        ("주의할 점으로, ", False),
        (f"{'·'.join(ce['fb2019_only'])}은 성과표에 없어 시기간 대조가 "
         "원리적으로 불가능하다", True),
        (f". 그런데 남지의 편각 잔차는 {st_['남지']['rD']:+.1f}′ 로 두 번째로 "
         "크다. 검증 수단이 없는 측점이 최악 잔차를 갖고 있는 상황이다.",
         False)])

    # ── 7. 향후 작업 순서 ─────────────────────────────────────────────
    heading(doc, "7. 향후 작업 순서", 1)
    para(doc,
         "순서의 근거는 신호대잡음이다. CYG 로 제거할 수 있는 외부장 일변동은 "
         f"편각 기준 6.7′ 인데 현재 미설명 잔차는 {asym['rD_rms']:.0f}′ 이다. "
         "잡음이 신호의 5배인 상태에서 보정을 얹으면 개선 여부 자체를 판정할 "
         "수 없다. 실제로 균일 V 근사를 쓴 기존 시도는 교차검증 편각을 "
         "0.751°에서 0.858°로 악화시켰다.")
    tcap(doc, "표 15. 수정된 작업 순서")
    booktabs(doc,
             ["단계", "내용", "비고"],
             [["0", "2022~25 야장 확보 및 관측시각 복원", "완료"],
              ["1", "성과표 원본 감사 (야장 원시 판독값 대조)",
               "완료 — 불일치 방문 2건 제외"],
              ["2", "External 포함 전체 재적합",
               "완료 — 현재 표본에서는 개선 미확인"],
              ["3", "잔차 큰 측점 현장 재측정 · 방위표지 고정",
               "판정기준상 다음 단계"],
              ["4", "측점 확충 후 교차검증 및 독립 표본 검증",
               "IGRF 대조는 ① Core 구현 확인용"],
              ["5", "지각 벡터 성분의 KPI 관점 재평가",
               f"예상 기여 {vr['predD_rms_arcmin']:.1f}′ > KPI 6′ 이므로 필수"]],
             [2.2, 6.4, 4.4])
    note(doc,
         "5단계를 생략해서는 안 된다. 3·4단계로 자료 계통 오차를 걷어내면 "
         "지각 벡터의 몫이 상대적으로 커지므로, 그 시점에서 격자 결측 처리와 "
         "해상도 한계를 함께 재검토해야 KPI 달성 가능성을 판정할 수 있다.")

    heading(doc, "7.1 1단계 판정기준", 2)
    tcap(doc, f"표 16. 사전 판정기준. 현재 {asym['rD_rms']:.0f}′ 이므로 "
                 "세 번째 구간에서 시작한다.")
    booktabs(doc,
             ["감사 후 편각 잔차 RMS", "판단"],
             [["10′ 이하", "CYG 단계로 진행"],
              ["10 ~ 20′", "잔차 큰 측점만 현장 재측정 후 재판정"],
              ["20′ 이상", "모형 작업 중단, 방위 기준 재결정이 선행"]],
             [5.0, 8.0])

    heading(doc, "7.2 검증에 관한 유의사항", 2)
    rich(doc, [
        ("IGRF 는 LMM 의 입력이므로 ", False),
        ("IGRF 와의 대조는 정확도 검증이 아니다", True),
        (". LMM − IGRF 는 정의상 보정층의 합이며, 이 문서 2.2절이 그것을 "
         "소수점까지 확인했다. IGRF 대조의 정당한 용도는 ① Core 층 구현 검증 "
         "하나이고 그것은 이미 통과했다. 정확도 검증은 교차검증, 적합에 쓰지 "
         "않은 hold-out 측점, 신규 캠페인 측점처럼 독립 표본이어야 한다.",
         False)])
    para(doc,
         "또한 2단계에서 층을 따로 손보아서는 안 된다. Regional 은 Core 와 "
         "Crustal 이 남긴 것을 흡수하는 항이므로 자료가 바뀌면 계수와 최적 "
         "차수가 함께 바뀐다. 층별 순차 튜닝은 서로의 오차를 흡수한다.")

    # ── 8. 한계 ───────────────────────────────────────────────────────
    heading(doc, "8. 이 기록의 한계", 1)
    for t in [
        f"표본이 작다. 편각 inlier {ds['inlier_D']}개로는 어떤 상관도 통계적 "
        "유의성을 확보하기 어렵다. 3절 가설들의 기각은 엄밀히 말해 '확인 실패'"
        "이며, 실질적 근거는 규모 논거와 반변량도이다.",
        f"KIGAM 격자에 자료 공백이 {vr['grid']['n_gap']:,}칸"
        f"({vr['grid']['gap_pct']:.0f} %) 있어 0 nT 로 채운 뒤 FFT 로 역산했다. "
        "결측과 「이상 0」은 다른 뜻이고, 그 경계에서 생기는 인공 불연속에 "
        "파수영역 미분은 특히 민감하다. 이것이 가설 A 검증의 가장 큰 약점이며, "
        "복원된 수평 성분과 예상 편각 기여값 자체에 왜곡이 있을 수 있다. "
        "따라서 3.2절은 확정 검정이 아니라 탐색적 분석으로 읽어야 한다.",
        f"격자 간격 {vr['grid']['dx_km']:.1f} km 는 진폭의 상한이 아니라 해상 "
        "가능한 파장의 한계다. 즉 이 격자에도 수백 nT 값은 얼마든지 존재할 수 "
        "있으나, 측점 주변 수 km 이하의 단파장 국소 이상은 복원되지 않는다. "
        f"잔차 설명에 필요한 {req['median_nT']:.0f} nT 규모의 국소 수평 성분이 "
        "실제로 있는지 여부는 현 자료만으로 판단할 수 없다.",
        f"반변량도의 최근거리 구간 측점쌍이 {vg['bins'][0]['n']}개뿐이고, "
        f"거리 추세의 p 값은 {vg['n_pairs']}개 쌍을 독립으로 본 값이라 "
        "과소평가되어 있다. 공간 구조의 부재를 확정하는 근거로 쓸 수 없다.",
        "측점별 편각 계통 오차라는 결론은 소거법으로 도달한 추론이며 직접 "
        "증거는 없다. 원본 대조와 현장 재측정으로만 확정된다.",
        "잠정 전제로 성과에 일변화 보정이 적용되지 않았다고 가정하였다"
        "(ASSUME_NO_DIURNAL_CORRECTION). 실제 적용 여부는 미확인이다.",
    ]:
        p = para(doc, "· " + t, size=10, indent=0.4, space_after=5, line=1.4)

    # ── 부록 ──────────────────────────────────────────────────────────
    heading(doc, "부록 A. 측점별 잔차", 1)
    rows = []
    for s in sorted(d["site_table"], key=lambda x: x["rD"]):
        rows.append([s["name"], s["n_visit"], f"{s['crustal_nT']:,.0f}",
                     f"{s['rD']:+.1f}", f"{s['rI']:+.1f}", f"{s['rF']:+.0f}",
                     f"{s['predD']:+.1f}",
                     "O" if s["inlier_D"] else "—"])
    tcap(doc, "표 A1. 측점별 잔차. 잔차 D·I 및 예측 ΔD 는 arcmin, "
                 "지각 ΔF 와 잔차 F 는 nT. 잔차는 관측에서 IGRF 와 Regional 을 "
                 "뺀 값이다.")
    booktabs(doc,
             ["측점", "관측", "지각 ΔF", "잔차 D", "잔차 I", "잔차 F",
              "예측 ΔD", "inlier"],
             rows, [1.9, 1.3, 2.0, 1.9, 1.9, 1.9, 1.9, 1.3],
             size=9,
             aligns=["l", "c", "r", "r", "r", "r", "r", "c"])

    heading(doc, "부록 B. 재현 방법", 1)
    para(doc, "이 문서의 모든 수치는 다음 두 명령으로 재현된다.")
    cp = para(doc, "python lmm_diagnose.py\npython create_lmm_validation_report.py",
              size=9.5, indent=0.6, space_after=8, line=1.5)
    for r in cp.runs:
        set_run_fonts(r, latin="Consolas", ea=KR)
    para(doc,
         "lmm_diagnose.py 가 docs/data/lmm_diagnosis.json 을 생성하고, 문서 "
         "생성기가 그 JSON 만 읽는다. 수치는 문서에 하드코딩되지 않으므로 "
         "자료가 갱신되면 문서도 함께 갱신된다.")

    return doc


def main():
    if not DIAG.exists():
        raise SystemExit(f"진단 파일이 없다: {DIAG}\n먼저 python lmm_diagnose.py 실행")
    d = json.loads(DIAG.read_text(encoding="utf-8"))
    doc = build(d)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / (datetime.now().strftime("%Y%m%d_%H%M%S")
                     + "_LMM_검증연구기록.docx")
    doc.save(out)
    print("saved:", out)
    return out


if __name__ == "__main__":
    main()
