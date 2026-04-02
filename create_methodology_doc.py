#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
대한민국 지구자기장 측정 입지 선정 시스템 — 기술 문서 생성
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="2E5090"):
    """스타일 적용된 표 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 헤더 행
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "맑은 고딕"
        set_cell_shading(cell, header_color)

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        row = table.rows[1 + r_idx]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "맑은 고딕"
            if c_idx >= 2 and len(headers) > 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r_idx % 2 == 1:
            for c in row.cells:
                set_cell_shading(c, "F2F6FA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # ── 기본 스타일 설정 ────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.3

    for level, size in [(0, 18), (1, 15), (2, 13)]:
        hs = doc.styles[f"Heading {level + 1}"]
        hs.font.name = "맑은 고딕"
        hs.font.size = Pt(size)
        hs.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)

    # ════════════════════════════════════════════════════════
    # 표지
    # ════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("대한민국 지구자기장 모델 구축을 위한\n측정 입지 선정 시스템")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
    run.bold = True
    run.font.name = "맑은 고딕"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("기술 문서 — 방법론·데이터·점수 산정 기준")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "맑은 고딕"

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("2026년 4월\n\n")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "맑은 고딕"
    run = info.add_run("도구: Python 3.12 + GeoPandas + Folium + OSM Overpass API")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = "맑은 고딕"

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 1. 개요
    # ════════════════════════════════════════════════════════
    doc.add_heading("1. 개요", level=1)
    doc.add_paragraph(
        "본 시스템은 대한민국 지구자기장 모델 구축을 위해 총자력계(proton magnetometer)를 "
        "이용한 지자기 측정 최적 입지를 자동으로 선정하는 도구이다. "
        "OpenStreetMap(OSM) 데이터를 기반으로 인공 간섭 요소를 식별·제외하고, "
        "다기준 평가를 통해 후보지의 우선순위를 산정한다."
    )
    doc.add_paragraph(
        "법적 근거로는 국립기상과학원 지자기관측 업무규정 제13조(도상선점) 및 "
        "제14조(선점실시)를 참조하되, 지도 출력물에서는 법조항 번호를 직접 표기하지 않는다."
    )

    # ════════════════════════════════════════════════════════
    # 2. 분석 범위 및 좌표계
    # ════════════════════════════════════════════════════════
    doc.add_heading("2. 분석 범위 및 좌표계", level=1)

    doc.add_heading("2.1 공간 범위", level=2)
    add_styled_table(doc,
        ["항목", "값"],
        [
            ["분석 대상", "대한민국 (남한) 육지부"],
            ["경계 Bounding Box", "경도 124.5°E ~ 129.6°E, 위도 33.0°N ~ 38.9°N"],
            ["경계 데이터 소스", "OSM Overpass API (admin_level=2, 대한민국)"],
            ["육지 마스크", "Natural Earth 50m 해상도 (제주도·서해 도서 포함)"],
            ["육지 마스크 버퍼", "2 km (연안 도서 보정, 해양 오염 최소화)"],
            ["해양 제외", "Overpass 경계 + naturalearth 50m 교차 → 영해 제거"],
        ],
        col_widths=[5, 12],
    )

    doc.add_heading("2.2 좌표 참조계", level=2)
    add_styled_table(doc,
        ["용도", "좌표계", "EPSG"],
        [
            ["공간 분석 (거리·버퍼)", "UTM Zone 52N", "EPSG:32652"],
            ["결과 출력 (지도·CSV)", "WGS84 경위도", "EPSG:4326"],
        ],
        col_widths=[6, 6, 4],
    )

    # ════════════════════════════════════════════════════════
    # 3. 후보 격자 생성
    # ════════════════════════════════════════════════════════
    doc.add_heading("3. 후보 격자 생성", level=1)
    add_styled_table(doc,
        ["항목", "값", "비고"],
        [
            ["격자 간격", "10 km", "UTM 좌표계 기준 정규 격자"],
            ["격자 유형", "정방형 포인트 격자", "np.arange(minx, maxx, 10000)"],
            ["1차 필터", "한국 경계 내부 (within)", "Overpass 경계 unary_union"],
            ["2차 필터", "육지 마스크 (within)", "naturalearth 50m + 2km 버퍼"],
            ["예상 후보점", "약 1,700 ~ 1,900개", "제외 전 기준"],
        ],
        col_widths=[4, 5, 8],
    )

    # ════════════════════════════════════════════════════════
    # 4. 제외 구역
    # ════════════════════════════════════════════════════════
    doc.add_heading("4. 제외 구역 (인공 간섭 요소)", level=1)
    doc.add_paragraph(
        "지자기 측정 시 인공적 자기 잡음을 발생시키는 시설물로부터 충분한 이격 거리를 확보하기 위해 "
        "아래 8개 항목을 제외 구역으로 설정한다. 모든 데이터는 OSM Overpass API에서 취득하며, "
        "JSON 형태로 로컬 캐시(data/ 디렉토리)하여 재쿼리를 방지한다."
    )

    add_styled_table(doc,
        ["번호", "제외 항목", "버퍼 반경", "OSM 태그 / 데이터", "적용 방식"],
        [
            ["[1]", "고압철탑·송전탑", "1.0 km",
             "power=line/tower/pole/station/\nsubstation/plant",
             "버퍼 폴리곤 합집합 (unary_union)"],
            ["[2]", "철도", "5.0 km",
             "railway=rail/subway/light_rail/\ntram/narrow_gauge",
             "LineString만 추출 → 버퍼"],
            ["[3]", "도시·주거", "폴리곤 직접 제외",
             "landuse=residential/commercial/\nindustrial/retail/construction",
             "폴리곤 합집합"],
            ["[4]", "파이프라인", "0.5 km",
             "man_made=pipeline",
             "버퍼 폴리곤 합집합"],
            ["[5]", "통신탑·기지국", "0.5 km",
             "man_made=mast/tower\n+ tower:type=communication",
             "버퍼 폴리곤 합집합"],
            ["[6]", "풍력발전기", "0.5 km",
             "generator:source=wind\nman_made=windmill",
             "버퍼 폴리곤 합집합"],
            ["[7]", "채석장·광산", "1.0 km",
             "landuse=quarry\nnatural=cave_entrance",
             "버퍼 폴리곤 합집합"],
            ["[8]", "자기이상도 고변화", "격자셀 직접 제외",
             "EMAG2 V3 데이터\n(gradient > 5 nT/km)",
             "Sobel 필터 → 고변화 셀 제외\n(EMAG2 배치 시 활성화)"],
        ],
        col_widths=[1.5, 3.5, 2.5, 4.5, 5],
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("참고: ")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = "맑은 고딕"
    run = p.add_run(
        "철도 버퍼 5.0 km은 직류/교류 구분이 불가하여 최대값(직류 기준) 적용. "
        "전력 인프라 1.0 km은 고압/일반 철탑 구분 불가로 고압 기준 적용. "
        "제외 구역 구축 시 비유효 geometry는 shapely.make_valid() 또는 buffer(0)으로 복구하며, "
        "대용량 폴리곤은 chunk_size=2000 단위 계층적 unary_union으로 메모리 최적화."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "맑은 고딕"

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 5. 입지 점수 산정 (핵심)
    # ════════════════════════════════════════════════════════
    doc.add_heading("5. 입지 점수 산정 기준", level=1)
    doc.add_paragraph(
        "제외 구역 필터링을 통과한 후보점에 대해 다기준 평가를 수행한다. "
        "총 8개 세부 지표(4개 평가 항목)로 구성되며, 현재 가용한 데이터로 산정 가능한 항목의 "
        "합산 점수를 100점 만점으로 정규화한다."
    )

    doc.add_heading("5.1 평가 항목 체계", level=2)
    add_styled_table(doc,
        ["평가 항목", "세부 지표", "배점", "가용", "산정 방식"],
        [
            ["공간 대표성", "① 격자 데이터 희소성", "25", "✅",
             "KDTree K=5 최근접 이웃 평균 거리 정규화"],
            ["", "② 지형적 대표성", "15", "❌",
             "DEM 기반 고도·경사도 표준편차 (미확보)"],
            ["환경 정온도", "③ 전력/철도 이격도", "15", "✅",
             "log(dist) 함수 기반 정규화"],
            ["", "④ 인구 밀집 이격도", "15", "✅",
             "log(dist) 함수 기반 정규화"],
            ["지질 안정성", "⑤ 자기 이상 균일도", "10", "⚠",
             "EMAG2 반경 20km 내 표준편차 역산"],
            ["", "⑥ 암상 적합성", "5", "❌",
             "지질도 기반 가중치 (미확보)"],
            ["운영 인프라", "⑦ 부지 지속성", "10", "❌",
             "국공유지 여부·개발 제한도 (미확보)"],
            ["", "⑧ 관리 접근성", "5", "❌",
             "도로망·통신망 거리 (미확보)"],
            ["합계", "", "100", "", ""],
        ],
        col_widths=[3, 4, 1.5, 1.2, 7.5],
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("✅ 산정 가능    ⚠ EMAG2 데이터 배치 시 활성화    ❌ 데이터 미확보 (향후 확장)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "맑은 고딕"

    # 5.2 세부 산정 방식
    doc.add_heading("5.2 세부 산정 방식", level=2)

    doc.add_heading("① 격자 데이터 희소성 (25점)", level=3)
    doc.add_paragraph(
        "후보점의 공간적 고립도를 측정한다. 주변에 다른 후보점이 적은(= 제외 구역에 의해 "
        "많이 제거된) 지역의 후보점일수록 높은 점수를 부여하여, 측정 네트워크의 공간적 대표성을 확보한다."
    )
    add_styled_table(doc,
        ["단계", "처리 내용"],
        [
            ["1", "scipy.spatial.KDTree로 전체 후보점의 좌표(UTM) 색인"],
            ["2", "각 후보점에서 가장 가까운 5개 이웃까지의 평균 거리 계산"],
            ["3", "평균 거리를 최대값으로 나누어 0~1 정규화"],
            ["4", "정규화 값 × 25 = 최종 점수 (0 ~ 25점)"],
        ],
        col_widths=[2, 15],
    )
    doc.add_paragraph(
        "해석: 주변 후보점이 멀리 분포할수록(= 고립도 높을수록) 높은 점수. "
        "이는 해당 지점이 측정 공백 지역에 위치하여 데이터 커버리지 확대에 기여함을 의미."
    )

    doc.add_heading("③ 전력/철도 이격도 (15점)", level=3)
    doc.add_paragraph(
        "전력 인프라와 철도로부터의 이격 거리가 클수록 전자기 간섭이 적어 "
        "측정 환경이 양호한 것으로 평가한다."
    )
    add_styled_table(doc,
        ["단계", "처리 내용"],
        [
            ["1", "각 후보점에서 전력 제외 구역(unary_union) 경계까지의 최소 거리 계산 (UTM, m)"],
            ["2", "각 후보점에서 철도 제외 구역 경계까지의 최소 거리 계산 (UTM, m)"],
            ["3", "거리에 log(1 + d) 변환 적용 (극단값 영향 완화)"],
            ["4", "각각 최대값으로 정규화 (0~1)"],
            ["5", "점수 = (전력 정규화 × 0.5 + 철도 정규화 × 0.5) × 15"],
        ],
        col_widths=[2, 15],
    )
    doc.add_paragraph(
        "참고: 제외 구역의 복잡 geometry는 simplify(500m)으로 간소화하여 연산 성능을 최적화. "
        "후보점은 이미 제외 구역 밖에 있으므로 최소 거리는 각 버퍼 반경 이상이 보장됨."
    )

    doc.add_heading("④ 인구 밀집 이격도 (15점)", level=3)
    doc.add_paragraph(
        "도시·주거 지역으로부터의 이격 거리가 클수록 인공 잡음이 적은 환경으로 평가한다."
    )
    add_styled_table(doc,
        ["단계", "처리 내용"],
        [
            ["1", "각 후보점에서 도시·주거 제외 구역 경계까지의 최소 거리 계산"],
            ["2", "log(1 + d) 변환 후 최대값 정규화"],
            ["3", "정규화 값 × 15 = 최종 점수 (0 ~ 15점)"],
        ],
        col_widths=[2, 15],
    )

    doc.add_heading("⑤ 자기 이상 균일도 (10점, 조건부)", level=3)
    doc.add_paragraph(
        "EMAG2 V3 데이터가 data/ 디렉토리에 배치된 경우에만 활성화된다. "
        "각 후보점 주변 반경 0.2°(약 20 km) 내의 자기이상(anomaly) 표준편차를 계산하여, "
        "균일한 자기장 환경일수록 높은 점수를 부여한다."
    )
    add_styled_table(doc,
        ["단계", "처리 내용"],
        [
            ["1", "후보점 좌표를 WGS84로 변환"],
            ["2", "EMAG2 격자에서 반경 0.2° 내 자기이상값 추출"],
            ["3", "해당 영역의 표준편차(std) 계산"],
            ["4", "점수 = (1 - std/max_std) × 10  →  낮은 std = 높은 점수"],
        ],
        col_widths=[2, 15],
    )

    doc.add_page_break()

    # 5.3 종합 점수 산출
    doc.add_heading("5.3 종합 점수 산출 및 등급 분류", level=2)
    doc.add_paragraph(
        "가용한 세부 지표의 점수를 합산하고, 가용 만점으로 나누어 100점 척도로 정규화한다."
    )
    add_styled_table(doc,
        ["구분", "포함 항목", "가용 만점", "비고"],
        [
            ["EMAG2 미배치", "①(25) + ③(15) + ④(15)", "55점", "55점 → 100점 정규화"],
            ["EMAG2 배치", "①(25) + ③(15) + ④(15) + ⑤(10)", "65점", "65점 → 100점 정규화"],
        ],
        col_widths=[3.5, 6, 3, 5],
    )

    doc.add_paragraph("")
    doc.add_paragraph("정규화 공식:")
    p = doc.add_paragraph()
    run = p.add_run("    종합 점수 = (가용 항목 합산 점수 / 가용 만점) × 100")
    run.bold = True
    run.font.name = "맑은 고딕"
    run.font.size = Pt(11)

    doc.add_heading("5.4 우선순위 등급 분류", level=2)
    doc.add_paragraph(
        "정규화된 종합 점수의 백분위수를 기준으로 3개 등급으로 분류한다."
    )
    add_styled_table(doc,
        ["등급", "조건", "의미", "색상"],
        [
            ["1등급 (최우선)", "상위 34% (점수 ≥ 66 percentile)", "측정 공백 해소·환경 최적", "🔴 빨강"],
            ["2등급 (우선)", "34% ~ 67% (33p ≤ 점수 < 66p)", "양호한 입지 조건", "🟠 주황"],
            ["3등급 (일반)", "하위 33% (점수 < 33 percentile)", "기본 조건 충족", "🟢 초록"],
        ],
        col_widths=[3.5, 5.5, 4.5, 3.5],
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "등급 분류 기준은 전체 후보점의 상대적 점수 분포(percentile)에 기반하므로, "
        "데이터셋이 변경되면 등급 경계값도 자동 조정된다. "
        "np.percentile(scores, [33, 66])으로 33번째·66번째 백분위 값을 산출하여 "
        "3분위 분류를 수행한다."
    )
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "맑은 고딕"

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 6. 데이터 소스
    # ════════════════════════════════════════════════════════
    doc.add_heading("6. 데이터 소스 및 캐싱", level=1)

    add_styled_table(doc,
        ["데이터", "소스", "포맷", "캐시 파일"],
        [
            ["대한민국 경계", "OSM Overpass API\n(admin_level=2)", "GeoJSON",
             "data/korea_boundary.geojson"],
            ["육지 마스크", "Natural Earth 50m\n(GitHub)", "GeoJSON",
             "data/naturalearth_kor_50m.geojson"],
            ["전력 인프라", "OSM Overpass API", "JSON", "data/power_infra.json"],
            ["철도", "OSM Overpass API", "JSON", "data/railways.json"],
            ["도시·주거", "OSM Overpass API", "JSON", "data/urban_residential.json"],
            ["파이프라인", "OSM Overpass API", "JSON", "data/pipelines.json"],
            ["통신탑·기지국", "OSM Overpass API", "JSON", "data/comm_towers.json"],
            ["풍력발전기", "OSM Overpass API", "JSON", "data/wind_turbines.json"],
            ["채석장·광산", "OSM Overpass API", "JSON", "data/quarries_mines.json"],
            ["자기이상도", "NOAA EMAG2 V3\n(수동 배치)", "CSV (8열)",
             "data/EMAG2_V3_20170530.csv"],
        ],
        col_widths=[3.5, 4, 2.5, 7],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Overpass API 쿼리는 3회 재시도(60초/120초 대기)를 지원하며, "
        "성공 시 JSON으로 로컬 캐시되어 이후 실행에서는 API 호출 없이 캐시를 사용한다. "
        "캐시 파일 삭제 시 다음 실행에서 재취득한다."
    )

    # ════════════════════════════════════════════════════════
    # 7. 출력물
    # ════════════════════════════════════════════════════════
    doc.add_heading("7. 출력물", level=1)

    doc.add_heading("7.1 대화형 HTML 지도", level=2)
    doc.add_paragraph("파일: output/geomag_site_selection.html")
    add_styled_table(doc,
        ["레이어", "표시 내용", "기본 표시"],
        [
            ["OpenStreetMap / 위성 이미지", "기본 타일 (전환 가능)", "O"],
            ["제외 구역 [1]~[8]", "반투명 폴리곤 + 범례 색상", "O"],
            ["후보지 1등급 (최우선)", "빨간 원형 마커 (반경 7)", "O"],
            ["후보지 2등급 (우선)", "주황 원형 마커 (반경 6)", "O"],
            ["후보지 3등급 (일반)", "초록 원형 마커 (반경 5)", "O"],
            ["격자점 전체 (참고)", "회색 소형 마커", "X"],
            ["1:50,000 도엽 격자", "15'×15' 셀 + 도엽명 라벨", "X"],
        ],
        col_widths=[5, 8, 3],
    )
    doc.add_paragraph(
        "도엽 라벨은 줌 레벨에 따라 동적으로 크기가 조절된다 "
        "(줌 9 미만: 숨김, 9~10: 7px, 10~11: 9px, 11~12: 11px, 12+: 13px)."
    )

    doc.add_heading("7.2 CSV 후보지 목록", level=2)
    doc.add_paragraph("파일: output/candidate_sites.csv")
    add_styled_table(doc,
        ["컬럼", "설명", "단위"],
        [
            ["site_id", "후보지 일련번호", "-"],
            ["lat", "위도 (WGS84)", "도(°)"],
            ["lon", "경도 (WGS84)", "도(°)"],
            ["priority", "우선순위 등급 (1=최우선, 2=우선, 3=일반)", "-"],
            ["score", "종합 점수 (100점 만점 정규화)", "점"],
            ["s1_희소성", "① 격자 데이터 희소성 점수", "/ 25"],
            ["s3_전력철도", "③ 전력/철도 이격도 점수", "/ 15"],
            ["s4_인구이격", "④ 인구 밀집 이격도 점수", "/ 15"],
            ["s5_자기균일", "⑤ 자기 이상 균일도 점수 (EMAG2)", "/ 10"],
            ["d_power_km", "전력 제외 구역까지 이격 거리", "km"],
            ["d_railway_km", "철도 제외 구역까지 이격 거리", "km"],
            ["d_urban_km", "도시·주거 제외 구역까지 이격 거리", "km"],
        ],
        col_widths=[3.5, 9, 3],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 8. 1:50,000 도엽 격자
    # ════════════════════════════════════════════════════════
    doc.add_heading("8. 1:50,000 지형도 도엽 격자", level=1)
    add_styled_table(doc,
        ["항목", "값"],
        [
            ["셀 크기", "15' × 15' (위경도 0.25° × 0.25°)"],
            ["격자 체계", "국제 지도 번호 체계 (NJ-52 등) 기반"],
            ["도엽명 출처", "국토지리정보원(NGII) 1:50,000 도엽명 현황(2014)"],
            ["수록 도엽 수", "약 160개 주요 도엽명 + 미수록 셀 좌표 기반 레이블"],
            ["클리핑", "한국 경계 폴리곤과 intersects하는 셀만 표시"],
            ["지도 레이어", "토글 가능 FeatureGroup (기본 비표시)"],
        ],
        col_widths=[4, 13],
    )

    # ════════════════════════════════════════════════════════
    # 9. 기술 스택
    # ════════════════════════════════════════════════════════
    doc.add_heading("9. 기술 스택 및 의존성", level=1)
    add_styled_table(doc,
        ["패키지", "용도", "비고"],
        [
            ["Python 3.12", "실행 환경", "Anaconda 배포판"],
            ["geopandas", "공간 데이터 처리·좌표 변환", "필수"],
            ["shapely", "geometry 연산 (buffer, union, distance)", "필수"],
            ["folium", "Leaflet.js 기반 대화형 HTML 지도", "필수"],
            ["numpy / pandas", "수치 계산·데이터프레임", "필수"],
            ["requests", "Overpass API HTTP 통신", "필수"],
            ["scipy", "KDTree (희소성 점수)", "선택 (미설치 시 중간값)"],
            ["pyproj / fiona", "좌표 변환 백엔드", "geopandas 의존"],
        ],
        col_widths=[3.5, 7, 6],
    )

    # ════════════════════════════════════════════════════════
    # 10. 향후 확장
    # ════════════════════════════════════════════════════════
    doc.add_heading("10. 향후 확장 가능 항목", level=1)
    doc.add_paragraph(
        "현재 미산정 항목(총 35점)은 해당 데이터 확보 시 코드 내 "
        "compute_priority() 함수에 추가 구현하여 활성화할 수 있다."
    )
    add_styled_table(doc,
        ["미산정 항목", "배점", "필요 데이터", "확보 방안"],
        [
            ["② 지형적 대표성", "15", "수치 표고 모델 (DEM)",
             "국토지리정보원 DEM 5m / SRTM 30m"],
            ["⑥ 암상 적합성", "5", "1:250,000 지질도",
             "한국지질자원연구원 (KIGAM) 지질도"],
            ["⑦ 부지 지속성", "10", "토지이용현황도 / 국공유지 DB",
             "국토교통부 토지이용규제정보서비스"],
            ["⑧ 관리 접근성", "5", "도로망 데이터",
             "OSM 도로 쿼리 또는 국가교통DB"],
        ],
        col_widths=[3.5, 1.5, 5, 7],
    )

    # ── 저장 ────────────────────────────────────────────────
    out_path = Path("C:/LG_gram_backup_users/LX/2026_geomag/output/site_selection_methodology.docx")
    doc.save(str(out_path))
    print(f"saved: {out_path}")


if __name__ == "__main__":
    main()
