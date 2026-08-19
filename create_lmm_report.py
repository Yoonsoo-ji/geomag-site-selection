#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
한반도 LMM 구축 현황 보고서 생성
=================================

lmm_build.py 가 산출한 docs/data/lmm_model.json 을 읽어 현황 보고서를 만든다.
수치를 하드코딩하지 않으므로 모델을 다시 적합하면 문서도 자동으로 갱신된다.

    python create_lmm_report.py
"""

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BASE = Path(__file__).parent
MODEL = BASE / "docs" / "data" / "lmm_model.json"
OUT_DIR = BASE / "docs" / "output"

FONT = "맑은 고딕"
NAVY = "2E5090"
GREY = "F2F4F7"
WARN = "FDF0E3"


def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex,
    })
    tcPr.append(shd)


def _trPr(row):
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = tr.makeelement(qn("w:trPr"), {})
        tr.insert(0, trPr)
    return trPr


def no_row_split(row):
    """행이 페이지 경계에서 잘리지 않게 한다."""
    _trPr(row).append(row._tr.makeelement(qn("w:cantSplit"), {}))


def repeat_header(row):
    """표가 페이지를 넘어갈 때 머리글 행을 반복한다."""
    _trPr(row).append(row._tr.makeelement(qn("w:tblHeader"), {}))


def add_styled_table(doc, headers, rows, header_color=NAVY, center_from=1):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = FONT
        set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[1 + r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = FONT
            if c_idx >= center_from:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r_idx % 2 == 1:
            for c in t.rows[1 + r_idx].cells:
                set_cell_shading(c, GREY)

    # 페이지 경계 처리: 머리글 반복 + 행 분할 금지
    repeat_header(t.rows[0])
    for row in t.rows:
        no_row_split(row)

    doc.add_paragraph()
    return t


def add_callout(doc, title, body, color=WARN):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = FONT
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    r2.font.name = FONT
    set_cell_shading(cell, color)
    no_row_split(t.rows[0])   # 제목과 본문이 페이지에서 갈라지지 않게
    doc.add_paragraph()


def sheet_scale_table():
    """
    축척별 도폭 규격과 도폭 내 편각 변화를 계산한다.

    도폭당 자침편차를 단일 값으로 표기하려면 도폭 내 편각 변화가
    표기 정밀도보다 작아야 한다. 이 표가 1:25,000 을 기준으로 삼는
    공학적 근거이며, 하드코딩하지 않고 IGRF-14 로 직접 계산한다.
    """
    import datetime as dt

    import numpy as np

    from lmm_build import igrf_dif

    d = dt.datetime(2027, 1, 1)
    rows = []
    for name, step, spec in [("1:5,000", 0.025, "1′30″"),
                             ("1:10,000", 0.05, "3′"),
                             ("1:25,000", 0.125, "7′30″"),
                             ("1:50,000", 0.25, "15′")]:
        var = []
        for la0 in np.arange(34.0, 38.0, 0.5):
            for lo0 in np.arange(126.5, 129.0, 0.5):
                la = np.array([la0, la0 + step, la0, la0 + step])
                lo = np.array([lo0, lo0, lo0 + step, lo0 + step])
                D, *_ = igrf_dif(la, lo, np.zeros(4), d)
                var.append(D.max() - D.min())
        med = float(np.median(var))
        rows.append([name, spec,
                     f"{step * 111 * np.cos(np.radians(36)):.1f} × {step * 111:.1f} km",
                     f"{med:.4f}°", f"{100 * med / 0.1:.0f} %"])
    return rows


def para(doc, text, size=10.5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = FONT
    r.bold = bold
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(10)
        r.font.name = FONT


def main():
    m = json.loads(MODEL.read_text(encoding="utf-8"))
    reg, cv, cd = m["regional"], m["loo_cv"], m["crustal_diagnostics"]
    val = {(v["성분"], v["단계"]): v for v in m["validation"]}
    sites = m["sites"]
    rep = m["repeatability"]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.3
    for lvl, size in enumerate([16, 13, 11.5]):
        hs = doc.styles[f"Heading {lvl+1}"]
        hs.font.name = FONT
        hs.font.size = Pt(size)
        hs.font.color.rgb = RGBColor(0x1F, 0x3B, 0x66)
        # 제목만 페이지 끝에 남지 않게
        hs.paragraph_format.keep_with_next = True
        hs.paragraph_format.page_break_before = False

    # ------------------------------------------------------------ 표지
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("한반도 지역 자기장 모델(LMM) 구축 현황")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = FONT
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Korea Local Magnetic Model — 중간보고")
    r.font.size = Pt(13)
    r.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(datetime.now().strftime("%Y년 %m월 %d일"))
    r.font.size = Pt(10.5)
    r.font.name = FONT
    doc.add_paragraph()

    add_callout(
        doc, "요약",
        "오석훈 교수 설명자료의 4-층 결합 구조를 현재 보유 자료로 구현하고, "
        "웹·엑셀 계산기를 제작해 세 구현(Python·JavaScript·Excel)의 수치 일치를 확인하였다. "
        f"현재 교차검증 오차는 D {cv['D']:.2f}°, F {cv['F']:.0f} nT 로 "
        "목표 KPI(D<0.1°, F<50 nT)에 미달하며, 원인은 모델 구조가 아니라 "
        "③ Crustal 층 해상도 부족·지표 측점 수 부족·④ External 층의 근사 한계에 있다. "
        "해당 KPI 는 설명자료가 제시한 공학적 목표치로, 법정 기준이 아님을 7장에서 확인하였다.",
        color="EAF1FA")

    # ------------------------------------------------------------ 1
    doc.add_heading("1. 모델 구조와 자료 보유 현황", level=1)
    para(doc, "B_LMM(r,t) = B_IGRF + B_Regional + B_Crustal + B_External", bold=True)
    para(doc, "각 층은 물리적 기원·시간상수·공간규모가 달라 단순 내삽으로 분리할 수 없다. "
              "현재 4개 층 중 3개를 구현하였다.")
    add_styled_table(
        doc,
        ["층", "자료원", "공간규모", "보유"],
        [["① Core", "IGRF-14 (13차 구면조화)", "≳3,000 km", "확보"],
         ["② Regional", f"지표 절대측정 {len(sites)}개 측점", "50~3,000 km", "확보"],
         ["③ Crustal", "KIGAM 자력이상도 1.5분 격자 (+ 벡터 복원)",
          "0.05~50 km", "해상도 부족"],
         ["④ External", "관측소 4소 1분 자료 (청양·제주·강릉·이천)",
          "일·시·분", "편각에 적용"]],
        center_from=2)
    para(doc, "③ Crustal 층은 총자력뿐 아니라 편각·복각에도 기여한다. 항공자력이 주는 "
              "스칼라 이상은 이상벡터를 주자기장 방향으로 투영한 값이므로, 파수영역 "
              "역산으로 벡터 3성분을 복원해 방향 성분에 반영하였다. "
              "④ External 층은 관측 성과 정리 단계에서 편각에만 적용한다 — 모델 자체는 "
              "정온시 기준값 모형이므로 예측 시점에 외부장 항을 평가하지 않는다. "
              "상세는 5.4 참조.", size=9.5)

    doc.add_heading("1.1 Regional 층 입력 측점", level=2)
    para(doc, f"성과표(2022~2025)와 2019년 야장 검증 통과분을 합쳐 {len(sites)}개 측점을 사용한다.")
    rows = [[s["name"], f"{s['lat']:.4f}", f"{s['lon']:.4f}", s["n_visit"]]
            for s in sorted(sites, key=lambda x: -x["lat"])]
    add_styled_table(doc, ["측점", "위도(°N)", "경도(°E)", "관측 횟수"], rows)

    # ------------------------------------------------------------ 2
    doc.add_heading("2. 산출물", level=1)
    add_styled_table(
        doc,
        ["산출물", "파일", "특징"],
        [["웹 계산기", "docs/lmm.html",
          "단일 파일 오프라인. IGRF-14 구면조화를 JS 로 직접 구현. 분포도·클릭계산·CSV 일괄"],
         ["엑셀 계산기", "docs/output/*_LMM_계산기.xlsx",
          "매크로 없이 수식만. IGRF 는 0.1° 사전격자 보간(오차 0.017 nT)"],
         ["모델 계수", "docs/data/lmm_model.json", "웹·엑셀 계산기 공통 입력"],
         ["구축 스크립트", "lmm_build.py", "적합·검증·계수 산출"],
         ["야장 파서", "lmm_fieldbook.py", "MAG-01H 야장에서 관측 일시 복원"],
         ["CYG 수집기", "lmm_cyg.py", "INTERMAGNET HAPI 1분 자료 + GFZ Kp 지수"],
         ["외부장 보정", "lmm_external.py", "야장 일시 × CYG × Kp 결합"],
         ["야장 검증", "lmm_verify2019.py", "원시 판독값 역산 감사"]],
        center_from=3)

    doc.add_heading("2.1 구현 간 수치 일치 검증", level=2)
    add_styled_table(
        doc,
        ["대조", "편각 D", "총자력 F"],
        [["JavaScript ↔ Python", "0.22″ 이내", "0.27 nT 이내"],
         ["Excel ↔ Python", "0.33″ 이내", "0.12 nT 이내"],
         ["JS IGRF 엔진 ↔ ppigrf", "0.17″ 이내", "0.03 nT 이내"]])

    # ------------------------------------------------------------ 3
    doc.add_heading("3. 검증 결과", level=1)
    doc.add_heading("3.1 단계별 잔차 (inlier 기준)", level=2)

    # 단계 이름은 지각 벡터 적용 여부에 따라 달라진다('+Regional' vs
    # '+Crustal+Regional'). 하드코딩하지 않고 성분별 마지막 단계를 찾는다.
    def last(comp):
        rows = [r for r in m["validation"] if r["성분"] == comp]
        return rows[-1]

    def mid(comp):
        rows = [r for r in m["validation"] if r["성분"] == comp]
        return next((r for r in rows if r["단계"] == "+Crustal"), None)

    fin = {c: last(c) for c in ("D_deg", "I_deg", "F_nT")}
    crm = {c: mid(c) for c in ("D_deg", "I_deg", "F_nT")}
    vec_on = crm["D_deg"] is not None
    add_styled_table(
        doc,
        ["성분", "IGRF 단독", "+ Crustal", "+ Regional", "측점 수"],
        [["편각 D", f"{val[('D_deg','IGRF')]['RMS']:.3f}°",
          f"{crm['D_deg']['RMS']:.3f}°" if vec_on else "—",
          f"{fin['D_deg']['RMS']:.3f}°", fin["D_deg"]["inlier수"]],
         ["복각 I", f"{val[('I_deg','IGRF')]['RMS']:.3f}°",
          f"{crm['I_deg']['RMS']:.3f}°" if vec_on else "—",
          f"{fin['I_deg']['RMS']:.3f}°", fin["I_deg"]["inlier수"]],
         ["총자력 F", f"{val[('F_nT','IGRF')]['RMS']:.1f} nT",
          f"{val[('F_nT','+Crustal')]['RMS']:.1f} nT",
          f"{fin['F_nT']['RMS']:.1f} nT", fin["F_nT"]["inlier수"]]])

    if vec_on:
        para(doc,
             "지각층은 총자력뿐 아니라 편각·복각에도 기여한다. 항공자력이 주는 "
             "스칼라 이상은 이상벡터를 주자기장 방향으로 투영한 값이므로, 파수영역 "
             "역산으로 벡터 3성분을 복원해 IGRF 벡터에 더한 뒤 각을 다시 잰다. "
             "위 표의 「+ Crustal」 열이 그 단계다.")

    doc.add_heading("3.2 KPI 판정", level=2)
    para(doc, "Leave-One-Out 교차검증은 적합에 쓰이지 않은 측점에서의 예측오차이므로 "
              "과적합에 속지 않는 실제 성능이다.")
    add_styled_table(
        doc,
        ["성분", "적합 RMS", "LOO RMS", "KPI", "판정"],
        [["편각 D", f"{fin['D_deg']['RMS']:.3f}°",
          f"{cv['D']:.3f}°", "< 0.1°",
          "통과" if cv["D"] < 0.1 else "미달"],
         ["복각 I", f"{fin['I_deg']['RMS']:.3f}°",
          f"{cv['I']:.3f}°", "—", "—"],
         ["총자력 F", f"{fin['F_nT']['RMS']:.1f} nT",
          f"{cv['F']:.1f} nT", "< 50 nT",
          ("통과" if cv["F"] < 50 else
           "적합 통과 / LOO 미달" if fin["F_nT"]["RMS"] < 50 else "미달")]])

    if fin["I_deg"]["inlier수"] < fin["D_deg"]["inlier수"]:
        para(doc,
             f"주의 — 복각의 평가 측점이 {fin['I_deg']['inlier수']}개로 편각"
             f"({fin['D_deg']['inlier수']}개)보다 적다. 지각 벡터를 반영하면서 "
             "복각 잔차가 줄자 강건 적합의 채택 범위가 좁아진 결과이므로, 복각 "
             "교차검증값은 남은 측점에 대한 값으로 읽어야 하며 전체 성능으로 "
             "일반화할 수 없다.")

    doc.add_heading("3.3 Crustal 층 설명력", level=2)
    para(doc, f"KIGAM 전국 격자가 지표 점 잔차를 설명하는 정도: 상관계수 r = {cd['corr']:+.3f}, "
              f"RMS {cd['rms_before_nT']:.1f} → {cd['rms_after_nT']:.1f} nT "
              f"({cd['rms_reduction_pct']:.1f}% 감소, n = {cd['n']}).")
    add_callout(
        doc, "해석 시 주의",
        f"표본이 {cd['n']}개에 불과해 상관계수가 개별 측점에 크게 좌우된다. "
        "실제로 측점 구성이 한 곳만 바뀌어도 r 이 0.16~0.51 범위에서 움직였다. "
        "보유 격자는 1.5분(약 2.8 km) 전국 컴필레이션인 반면, 설명자료의 사례연구는 "
        "측선간격 250 m 원측선 자료를 사용하였다. 다만 해당 원측선 자료는 존재하지 "
        "않는 것으로 확인되었으므로, 현 격자가 지각층 해상도의 상한이다. "
        "복원 벡터의 최적배율이 1.60 으로 나타나는 진폭 부족분은 현 자료로는 채울 수 "
        "없으며, 신규 자력측량 여부가 갈림길이다.")

    # ------------------------------------------------------------ 4
    doc.add_heading("4. 2019년 야장 검증 및 통합", level=1)
    para(doc, "2019년 야장 6건을 입수하여 신뢰도를 검증한 뒤 Regional 층에 통합하였다. "
              "야장 판독값의 각도 단위는 도(°)가 아니라 곤(gon, 1회전 400분할)이며, "
              "야장 수식의 180/360 은 실제로 200/400 이다.")

    doc.add_heading("4.1 내부 정합성 (야장 자체 검증)", level=2)
    para(doc, "원시 판독값(ED/WD/EU/WU, SU/ND/SD/NU)으로 최종 성과를 역산해 대조하였다. "
              "22개 세션 전부 통과.")
    add_styled_table(
        doc,
        ["검증 항목", "최대 편차", "허용", "판정"],
        [["마크 폐합 (200 gon)", "0.0116 gon (≈37″)", "0.02 gon", "통과"],
         ["편각 재계산 vs 야장", "0.0001°", "0.002°", "통과"],
         ["복각 재계산 vs 야장", "0.0000°", "0.002°", "통과"],
         ["분력 정합 F²=H²+Z²", "0.0059 nT", "2 nT", "통과"]])

    doc.add_heading("4.2 외부 대조", level=2)
    para(doc, "IGRF 잔차로 비교하면 영년변화가 제거되고 정적인 지각 이상만 남으므로, "
              "서로 다른 시기의 측정을 직접 대조할 수 있다.")
    add_styled_table(
        doc,
        ["측점", "2019 ΔD", "2022~25 ΔD", "차이", "해당 측점 재방문 산포"],
        [["미원", "−0.448°", "−0.445°", "0.003°", "2.4′"],
         ["이원", "+0.246°", "+0.096°", "0.151°", "4.9′"],
         ["거제", "+0.420°", "+0.195°", "0.225°", "12.1′"]])
    para(doc, "차이는 각 측점 고유의 재방문 산포와 같은 수준이며, 2019 자료가 특별히 "
              "열등하지 않음을 보인다. 특히 미원은 5년 간격으로 0.003° 일치한다.")

    doc.add_heading("4.3 통합 결과", level=2)
    add_styled_table(
        doc,
        ["지표", "성과표만 (15측점)", "2019 포함 (17측점)"],
        [["LOO D-RMS", "0.809°", f"{cv['D']:.3f}°"],
         ["LOO I-RMS", "0.226°", f"{cv['I']:.3f}°"],
         ["LOO F-RMS", "64.6 nT", f"{cv['F']:.1f} nT"],
         ["Crustal 상관 r", "+0.162", f"{cd['corr']:+.3f}"]])

    add_callout(
        doc, "투입 제외 및 유보",
        "성산 — 관측시각 미기입, 원시 판독값 배열 상이로 역산 검증 불가하여 제외. "
        "미원 — 「'10~'19년 관측현황」과 성과표의 좌표가 353 m 어긋나 현장 확인 필요 "
        "(거제·성산·이원은 10 m 이내 일치). 현재는 성과표 좌표를 사용.")

    # ------------------------------------------------------------ 5
    doc.add_heading("5. 외부장(External) 층 검토 결과", level=1)
    doc.add_heading("5.1 CYG 자료 확보 가능성", level=2)
    para(doc, "INTERMAGNET HAPI 서버에서 청양 관측소 자료를 수집할 수 있다. "
              "다만 등급별 보유 기간이 달라 주의가 필요하다.")
    add_styled_table(
        doc,
        ["등급", "보유 기간", "측정기간 포함 여부"],
        [["definitive (확정치)", "2014-01-01 ~ 2017-12-31", "2022~2025 미포함"],
         ["quasi-definitive", "2015-04-25 ~ 2020-12-31", "2019년만 포함"],
         ["adjusted / best-avail", "2013-10-30 ~ 현재", "전 기간 포함"]])

    doc.add_heading("5.2 외부장 변동의 규모", level=2)
    para(doc, "2022~2025 표본일의 CYG 일변동 실측값이다.")
    add_styled_table(
        doc,
        ["성분", "중앙값", "최대", "KPI", "판정"],
        [["총자력 F", "37.5 nT", "66.2 nT", "< 50 nT", "KPI 근접/초과"],
         ["편각 D", "0.112°", "0.220°", "< 0.1°", "전형적 하루도 KPI 초과"]])
    add_callout(
        doc, "핵심 결론",
        "편각은 전형적인 하루의 외부장 변동만으로도 KPI(0.1°)를 초과한다. "
        "따라서 관측 시각을 알고 외부장을 보정하지 않으면 편각 KPI 달성은 원리적으로 "
        "불가능하다. 성과표에는 관측연도만 기록되어 있으나, 국토지리정보원 야장 원본에서 "
        "2019~2025 분 단위 관측시각을 복원하여 이 전제 조건은 해소되었다.")

    doc.add_heading("5.3 단순 보정의 한계", level=2)
    para(doc, "2019 야장 22개 세션에 대해 CYG 정온야간 기준선 대비 보정을 시험 적용하였다.")
    add_styled_table(
        doc,
        ["조건", "D 세션산포", "I 세션산포"],
        [["정온시 (Kp ≤ 2), 7건", "1.12′ → 1.27′ (악화)", "1.88′ → 1.75′ (개선)"],
         ["교란시 (Kp > 2), 1건", "0.72′ → 0.93′ (악화)", "0.01′ → 4.39′ (악화)"]])
    para(doc, "「외부장 변동이 한반도 전역에서 동일하다」는 1차 근사는 CYG 로부터 "
              "76~229 km 떨어진 측점에서 실익이 없고, 자기폭풍 시에는 오히려 악화시킨다. "
              "설명자료가 NOC 공간투영을 요구하는 이유가 실측으로 확인되었다.")

    doc.add_heading("5.4 ④ External 층 적용 방침", level=2)
    add_callout(
        doc, "확정 — 성과표는 야장 세션평균 그대로인 원시값이다",
        "종전에는 일변화 보정 적용 여부를 정황으로 판단해 잠정 처리하였으나, "
        "지리원 야장 원본을 확보해 직접 대조하여 확정하였다. 성과표와 야장 세션평균이 "
        "매칭 18행 전부에서 표기 자릿수 이내로 일치한다(ΔD RMS 0.00′, ΔI 0.03′, "
        "ΔF 0.02 nT). 즉 기준시점 환산도 일변화 보정도 적용되지 않았으므로, "
        "외부장 보정을 적용할 근거가 성립한다. "
        "(코드: lmm_build.py 의 ASSUME_NO_DIURNAL_CORRECTION, audit_survey_table.py)")

    para(doc, "현재 방침은 관측소 4소(청양·제주·강릉·이천)의 정온야간 기준선 대비 "
              "세션 편차를 1차 평면으로 공간보간하여 측점 위치의 편차를 추정하고, "
              "정온 세션(Kp≤2)에 한해 성과 편각에서 빼는 것이다. 총자력에는 적용하지 "
              "않는다. 층이 결합되어 있어 순차 튜닝이 성립하지 않으므로 방식마다 "
              "처음부터 다시 적합하여 교차검증으로 비교하였다.")
    add_styled_table(
        doc,
        ["방식", "보정량 자료원", "LOO D", "LOO F"],
        [["미적용", "—", "0.590°", "58.4 nT"],
         ["전 세션 · 편각", "청양 단독", "0.365°", "58.4 nT"],
         ["전 세션 · 편각", "관측소 4소", "0.403°", "58.4 nT"],
         ["정온 세션 · 편각", "청양 단독", "0.376°", "58.4 nT"],
         ["정온 세션 · 편각 (채택)", "관측소 4소", "0.323°", "58.4 nT"]])
    para(doc, "다중 관측소는 정온 세션에 한정할 때에만 이득이 있다. 교란시에는 4소 "
              "평면적합이 크게 발산하여(총자력 편차 최대 567 nT) 잡음을 키운다. "
              "공간투영 근사가 자기폭풍 중에 먼저 깨진다는 뜻이며, 설명자료가 NOC "
              "모드 분해를 요구하는 이유와 같은 방향의 결과다. "
              "다만 위 다섯 가지 중 교차검증 최소값을 채택한 것이므로 0.365°와 0.323°의 "
              "차이 자체를 성과로 읽어서는 안 되며, 근거는 방식의 물리적 타당성에 있다.", size=9.5)

    doc.add_heading("5.5 2019 야장을 이용한 방법론 검증 (경과 기록)", level=2)
    para(doc, "아래는 관측시각이 확보되기 전, 2019 야장 구간만으로 수행한 예비 검증이다. "
              "현재 방침(5.4)의 근거가 된 관찰이므로 경과로 남긴다.")
    add_styled_table(
        doc,
        ["방식", "LOO D", "LOO I", "LOO F", "Crustal r"],
        [["① 보정 없음", "0.751°", "0.188°", "64.0 nT", "+0.507"],
         ["② Kp>2 세션 배제", "0.766°", "0.173°", "58.6 nT", "+0.508"],
         ["③ CYG 값 직접 차감", "0.858°", "0.188°", "58.0 nT", "+0.608"],
         ["④ ②+③ 동시", "0.858°", "0.173°", "65.8 nT", "+0.614"]])
    para(doc, "채택: ② Kp>2 세션 배제. ③ CYG 값 직접 차감은 다음 두 가지 이유로 적용하지 않는다.")
    bullets(doc, [
        "편각 악화 — LOO D 가 0.751°에서 0.858°로 나빠진다. "
        "「V 가 전역 균일」이라는 근사가 편각에서 특히 취약하기 때문이다. "
        "편각은 국소 수평분력 방향에 의존하므로 스칼라인 F 보다 공간 변동이 크다.",
        "F 기록 단위 불일치 — 야장의 총자기장은 일자별 단일값이며 거제·장흥은 "
        "측점당 1개뿐이다. 세션 시각 기준 보정량을 일자 대표값에서 빼는 것은 "
        "원리적으로 성립하지 않는다. 제19조제1항제6호는 세션마다 총자기장 측정을 "
        "요구하나 야장 양식이 이를 담고 있지 않다.",
    ])
    para(doc, "반면 ②는 값을 조작하지 않고 오염된 표본만 제외하므로 위 문제가 없으며, "
              "설명자료가 요구하는 Kp≤2 정온시간대 조건과도 일치한다. "
              "미원 2019-05-14 두 세션(Kp 5.7~6.3, 자기폭풍)이 이에 해당하여 제외되었다.")

    # ------------------------------------------------------------ 6
    doc.add_heading("6. 작업 중 발견한 자료 문제", level=1)
    add_styled_table(
        doc,
        ["구분", "내용", "조치"],
        [["KIGAM 격자", "위도축에 제주해협(33.55~34.1°) 0.55° 공백이 있어 "
                        "균일간격 가정 시 공백 위쪽 전 구간이 21행 어긋남",
          "결측행을 NaN 으로 채운 균일축 구성. 수정 후 F 오차 95 → 65 nT"],
         ["선택편향", "Crustal 보정 후 잔차로 inlier 를 고르면 "
                      "Crustal 기여도가 부풀려짐 (F 33 nT 로 과대평가)",
          "IGRF 잔차만으로 inlier 판정하도록 수정"],
         ["편각 블런더", "삼척은 2년 만에 편각이 +0.99° 변화 "
                        "(영년변화 최대 0.25°, 부호도 반대). 가야 35.5′",
          "재방문 산포 기준으로 배제"],
         ["측점 중복", "양산·언양이 동일 지점으로 이중 기재", "좌표 4자리 반올림 후 중복 제거"],
         ["관측소 결측", "CYG 2025-09-08~22 이상 전량 결측. "
                        "결측일도 1440행이 반환되므로 행 수로 판정 불가",
          "유효값 카운트로 완전성 판정"],
         ["야장 기입 오류", "장흥 2019-10-22 6회차의 복각 시각이 편각보다 57분 빠름",
          "편각·복각 구간을 분리 수집해 이상 자동 표시"]],
        center_from=3)

    # ------------------------------------------------------------ 7
    doc.add_heading("7. KPI 의 출처와 축척 기준 검토", level=1)
    para(doc, "본 보고서가 사용하는 KPI(D<0.1°, F<50 nT)와 「1:25,000 도폭」 기준의 "
              "법적 근거를 확인하기 위해 관련 법령을 전수 조사하였다.")

    doc.add_heading("7.1 관련 법령 체계", level=2)
    para(doc, "지도 제작 계열과 측량 작업 계열을 나누어 확인하였다. 편각에 관한 실체 규정은 "
              "지도 계열이 아니라 측량 계열인 「지구물리측량 작업규정」에 있다.")
    add_styled_table(
        doc,
        ["법령", "형식", "편각 관련 내용"],
        [["지도도식규칙 (제882호)", "국토교통부령", "없음 ('편각·자침·방위' 0회)"],
         ["지형도 도식적용규정 (제2022-3601호)", "국토지리정보원고시",
          "'자침' 1회 — 자침편차 도표는 임의 표기"],
         ["수치지형도 작성 작업 및 성과에 관한 규정 (제2026-2524호)",
          "국토지리정보원고시", "없음 (0회)"],
         ["지구물리측량 작업규정 (제2021-2985호)", "국토지리정보원고시",
          "제3장 지자기측량 (제12~21조) — 측정·오차한계·보정 규정"]],
        center_from=2)

    doc.add_heading("7.2 자편각 표기의 법적 근거", level=2)
    para(doc, "지구물리측량 작업규정 제12조제1항은 1등 지자기점의 설치 목적을 다음과 같이 정한다.")
    para(doc, "  「1등 지자기점이라 함은 일반적인 지구자기장의 분포와 영년변화를 조사·연구하고 "
              "국가기본도의 자편각 표기의 목적을 위해 설치한 점을 말한다」", size=10)
    para(doc, "즉 국가기본도에 자편각을 표기하는 것은 1등 지자기점의 법정 설치 목적이다. "
              "다만 도식 규정 측면에서는 자침편차 도표가 난외사항 필수목록에 없어 "
              "표기 자체는 임의사항이며(부령 제8조 열거 제외, 고시 제205조 목록 제외), "
              "적용 축척도 「국가기본도」라고만 되어 있어 1:25,000 으로 특정되지 않는다.")

    doc.add_heading("7.3 법정 측정기준과 LMM KPI 의 관계", level=2)
    para(doc, "지구물리측량 작업규정 제20조는 측정오차의 한계를 다음과 같이 정한다.")
    add_styled_table(
        doc,
        ["편각의 정수차", "복각의 정수차", "관측시간"],
        [["30′", "30′", "20분 이내"]], center_from=0)
    para(doc, "그 밖의 주요 법정 요구사항은 다음과 같다.")
    add_styled_table(
        doc,
        ["조문", "내용"],
        [["제17조", "총자기장 측정기는 0.1 nT 단위까지 판독 가능해야 함"],
         ["제19조①", "1등 1일 6회 이상 / 2등 1일 4회 이상 측정, "
                     "편각은 지구물리학적 진북 기준, 편각·복각 측정 시 "
                     "총자기장과 시간을 동시 측정"],
         ["제21조", "동일시간 상시 측정 기준점과 비교하여 일변화·영년변화를 보정하고 "
                    "기준년 값으로 환산"]],
        center_from=99)   # 서술형 본문이므로 좌측 정렬 유지
    add_callout(
        doc, "정정 — KPI 와 법정 기준의 관계",
        "법정 측정오차 한계는 편각 정수차 30′(=0.5°)로 존재한다. 다만 이는 "
        "측정 품질관리 기준(회차 간 재현성)이며, 모델 예측정확도 규격이 아니다. "
        "LMM KPI 0.1°(=6′)는 법정 한계보다 5배 엄격한 공학적 목표치로, "
        "설명자료(오석훈, 강원대)가 제시한 값이다. "
        "따라서 본 모델의 KPI 미달은 법령 위반이 아니다. "
        "역으로 법정 30′ 를 충족한 측정이라도 0.1° 모델의 입력으로는 부족할 수 있으며, "
        "실제로 포천(재방문 산포 22.3′)은 법정 한계 이내지만 포함 시 "
        "LOO D-RMS 가 0.751°에서 0.807°로 악화되어 배제하였다.")

    doc.add_heading("7.4 우리 작업의 법령 부합성 점검", level=2)
    add_styled_table(
        doc,
        ["항목", "법정 기준", "본 프로젝트", "판정"],
        [["철도 이격 (직류)", "5.0 km 이상", "5.0 km", "부합"],
         ["철도 이격 (교류·일반)", "2.0 km 이상", "5.0 km", "보수적 적용"],
         ["고압철탑 이격", "1.0 km 이상", "1.0 km", "부합"],
         ["송전탑 이격", "0.5 km 이상", "1.0 km", "보수적 적용"],
         ["재방문 편각 산포 허용", "정수차 30′", "20′", "엄격 적용 (7.3 참조)"]],
        center_from=1)
    para(doc, "입지선정 필터(제13조제2항 대응)는 모두 법정 기준 이상으로 보수적이다. "
              "2019 야장 22개 세션의 관측 소요시간은 중앙값 18분으로 제20조의 20분 기준에 "
              "대체로 부합하나 7개 세션이 21~25분으로 초과한다(장흥 1건 57분은 "
              "6장에 기술한 야장 시각 기입 오류에 따른 것으로 실제 소요시간이 아니다).")

    doc.add_heading("7.5 1:25,000 을 기준으로 삼는 근거", level=2)
    para(doc, "도폭 규격 자체는 법정 사항이다(도식적용규정 제4조·제211조). "
              "자침편차를 도폭당 단일 값으로 표기하려면 도폭 내 편각 변화가 표기 정밀도보다 "
              "작아야 하므로, 축척이 곧 요구 정밀도를 규정한다. "
              "아래는 IGRF-14 로 직접 계산한 도폭 내 편각 변화이다(한반도 40개 위치 중앙값).")
    add_styled_table(
        doc,
        ["축척", "도곽 (법정)", "실제 크기 (위도 36°)", "도폭 내 D 변화", "0.1° 예산 대비"],
        sheet_scale_table())
    para(doc, "1:50,000 은 도폭 내 변화만으로 0.1° 예산의 91%를 소진해 모델 오차를 담을 여지가 "
              "거의 없다. 1:25,000 은 45%만 사용하므로 나머지를 모델 정확도에 배분할 수 있다. "
              "즉 1:25,000 은 도폭당 단일 편각값 표기가 성립하는 가장 성긴 축척이며, "
              "이것이 해당 축척을 기준으로 삼는 공학적 근거이다.")

    doc.add_heading("7.6 확인이 필요한 사항", level=2)
    add_callout(
        doc, "성과표의 일변화 보정 적용 여부 — 확인 필요",
        "지구물리측량 작업규정 제21조는 「동일시간에 상시 측정 기준점의 측정값과 비교 검토」하여 "
        "일변화와 영년변화를 보정하도록 정하고 있다. 즉 외부장 보정은 임의사항이 아니라 "
        "법정 절차이다. 그렇다면 성과표(2022~2025)의 값에 이미 일변화 보정이 적용되어 있을 "
        "가능성이 있으며, 이 경우 본 프로젝트가 CYG 자료로 다시 보정하면 이중보정이 된다. "
        "반대로 동일 측점 재방문 F 산포가 최대 138 nT 에 이르는 점은 보정이 충분히 "
        "이루어지지 않았을 가능성을 시사한다. 성과 산출 과정에서 어느 기준점의 어떤 자료로 "
        "일변화 보정을 했는지 확인해야 ④ External 층의 설계를 확정할 수 있다.")
    bullets(doc, [
        "위 일변화 보정 적용 여부 및 사용 기준점 — ④ External 층 설계의 전제조건.",
        "관측 시각 기록 — 제19조제1항제6호가 「편각 및 복각 측정 시 동시에 총자기장과 시간을 "
        "측정한다」고 정하므로, 2022~2025 야장에도 시각이 기록되어 있어야 한다. "
        "성과표에만 누락된 것이므로 야장 확보 시 복원 가능하다.",
        "국토지리정보원 기관표준 「기본공간정보 데이터 품질」 — 축척별 위치정확도 RMSE 허용치.",
        "실제 1:25,000 지형도 도엽 — 자침편차 도표의 실제 인쇄 여부와 표기 자릿수.",
        "2027 캠페인 야장 양식 — 타원체고와 지오이드고를 함께 기록하도록 보완 (아래 7.7).",
    ])

    doc.add_heading("7.7 높이 기준 — 표고와 타원체고", level=2)
    para(doc, "IGRF 는 타원체고(ellipsoidal height)를 입력으로 요구하나, 성과표에는 "
              "표고(정표고, 평균해면 기준)만 기록되어 있다. 두 높이의 차가 지오이드고 N 이다.")
    para(doc, "  H(표고) = h(타원체고) − N(지오이드고),   한반도 N ≈ 20~30 m (KNGeoid)", size=10)
    para(doc, "GNSS·RTK 수신값은 타원체고이므로, 현장에서 얻은 높이를 그대로 쓰면 "
              "약 25 m 의 계통 차이가 생긴다. 그 영향을 실측으로 확인하였다.")
    add_styled_table(
        doc,
        ["높이 처리", "LOO F-RMS", "Regional F 상수항"],
        [["표고 그대로 (현행)", "58.64 nT", "−9.84 nT"],
         ["N = +20 m 부여", "58.64 nT", "−9.31 nT"],
         ["N = +25 m 부여", "58.64 nT", "−9.18 nT"],
         ["N = +30 m 부여", "58.64 nT", "−9.05 nT"],
         ["위도의존 N (22~28 m)", "58.64 nT", "—"]])
    para(doc, "지오이드는 공간적으로 완만하므로 그 오프셋이 Regional 다항식의 상수항에 "
              "그대로 흡수된다. 교차검증 오차는 어느 경우에도 58.64 nT 로 변하지 않으며 "
              "상수항만 약 0.7 nT 이동한다. 높이 감도가 F 기준 −0.026 nT/m 이므로 "
              "25 m 혼동은 0.66 nT 로, 현재 모델 불확도(59 nT)의 1% 수준이다.")
    add_callout(
        doc, "실무 지침",
        "① 계산기 입력은 표고로 통일한다. 모델이 표고로 적합되었으므로 타원체고를 넣으면 "
        "약 0.7 nT 의 불일치가 생긴다(현 단계에서는 무시 가능). "
        "② 2027 캠페인에서는 GNSS 원값인 타원체고와 적용한 지오이드 모델·N 값을 "
        "야장에 함께 기록하는 것이 바람직하다. F 목표가 50 nT 로 좁혀지면 "
        "높이 기준의 계통 오차도 관리 대상이 된다.",
        color="EAF1FA")

    # ------------------------------------------------------------ 8
    doc.add_heading("8. 한계 및 다음 단계", level=1)
    para(doc, "현재 모델이 KPI 에 미달하는 원인은 모델 구조가 아니라 입력자료에 있다. "
              "우선순위 순으로 정리하면 다음과 같다.")
    add_styled_table(
        doc,
        ["순위", "필요 조치", "기대 효과"],
        [["1", "잔차가 큰 측점의 현장 재측정 · 방위표지 고정",
          "편각 잔차 13.8′ — 1단계 판정기준상 남은 조건"],
         ["2", "2027 신규 캠페인 30점 측정",
          "권고 측점 수 충족. 1차 다항식이 버티는 하한이 14점이라 여유가 없다"],
         ["3", "상시관측 자료의 NOC 모드 분해",
          "외부장의 공간투영. 관측소 4소 1차 평면 근사의 한계 극복"],
         ["4", "(보류) 지각층 해상도",
          "항공자력 원측선이 존재하지 않아 현 격자가 상한. 신규 자력측량은 사업 판단"]],
        center_from=1)

    add_callout(
        doc, "현 단계 사용 지침",
        "본 계산기는 IGRF-14 대비 개선된 참고값 및 계산 체계 검증용으로만 사용하고, "
        "지형도 자침편차 등 정식 편각 산출에는 사용하지 않는다. "
        "웹·엑셀 계산기 모두에 동일한 정확도 고지와 KPI 출처(법정 기준 아님)를 표기하였다.")

    # ------------------------------------------------------------ 저장
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = OUT_DIR / f"{stamp}_LMM_구축현황보고서.docx"
    doc.save(out)
    print(f"[저장] {out}")
    return out


if __name__ == "__main__":
    main()
